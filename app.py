from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, session
import weaviate
import os
from PyPDF2 import PdfReader
import textwrap
from werkzeug.utils import secure_filename
import logging
from datetime import datetime, timedelta
from dotenv import load_dotenv
import requests
import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import io
from bs4 import BeautifulSoup
from document_crew import DocumentCrew
import traceback
from document_review_crew import review_document
import tempfile
from functools import wraps
from models.user import User
import atexit
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Inches
from models.client import Client
from models.case import Case
from models.appointment import Appointment
from models.initial_inquiry import InitialInquiry
from models.attorney import Attorney
from bson import ObjectId  # Add this import at the top if not present
import uuid
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,  # Change to DEBUG level
    format='%(asctime)s - %(levelname)s - %(name)s - %(message)s',  # Add logger name
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('app.log', mode='a')  # Append mode
    ]
)
logger = logging.getLogger(__name__)

# Disable MongoDB heartbeat debug messages
logging.getLogger('pymongo.monitoring').setLevel(logging.WARNING)
logging.getLogger('pymongo.connection').setLevel(logging.WARNING)
logging.getLogger('pymongo.topology').setLevel(logging.WARNING)

# Add more detailed logging for Weaviate operations
logging.getLogger('weaviate').setLevel(logging.DEBUG)

# Add logging for Flask
logging.getLogger('werkzeug').setLevel(logging.DEBUG)

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Initialize MongoDB collections
from models.user import User
db = User.db
clients_collection = db.clients
cases_collection = db.cases
documents_collection = db.documents
appointments_collection = db.appointments
initial_inquiries_collection = db.initial_inquiries

# Add nl2br filter
@app.template_filter('nl2br')
def nl2br_filter(s):
    if not s:
        return ''
    return s.replace('\n', '<br>')

# Register cleanup function
@atexit.register
def cleanup():
    """Cleanup function to close MongoDB connection when the application exits"""
    try:
        if hasattr(User, 'client'):
            User.client.close()
            logger.info("MongoDB connection closed")
    except Exception as e:
        logger.error(f"Error closing MongoDB connection: {str(e)}")

# Login required decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            flash('Please login first')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Authentication routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        if User.check_password(username, password):
            session['username'] = username
            flash('Successfully logged in!')
            return redirect(url_for('document_creation'))
        else:
            flash('Invalid username or password')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    flash('Registration is currently disabled')
    return redirect(url_for('login'))

@app.route('/logout')
def logout():
    session.pop('username', None)
    flash('Successfully logged out!')
    return redirect(url_for('login'))

# Protect existing routes
@app.route('/')
@login_required
def index():
    return redirect(url_for('document_creation'))

@app.route('/index')
@login_required
def search_page():
    try:
        if client is None:
            logger.error("Weaviate client is not initialized")
            return render_template('search.html', 
                                documents=[], 
                                error="Weaviate connection is not available. Please check your configuration.",
                                current_page=1,
                                total_pages=1,
                                per_page=10,
                                query='',
                                doc_type='',
                                doc_types=[],
                                all_documents=[])

        # Get query parameters
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 10))
        doc_type = request.args.get('doc_type', '')
        
        logger.info(f"Fetching documents with type filter: {doc_type}")
        
        # Prepare type filter
        filter_clause = None
        if doc_type:
            doc_type = doc_type.replace('%20', ' ')
            logger.info(f"Searching for documents with type: {doc_type}")
            
            filter_clause = {
                        "path": ["doc_type"],
                        "operator": "Equal",
                "valueString": doc_type
            }

        # Fetch all documents using offset-based pagination
        offset = 0
        batch_size = 3000
        all_results = []

        while True:
            query = (
            client.query
                .get("Document", ["document_name", "doc_type"])
            .with_additional("id")
                .with_limit(batch_size)
                .with_offset(offset)
            )

            if filter_clause:
                query = query.with_where(filter_clause)

            result = query.do()
            documents = result.get("data", {}).get("Get", {}).get("Document", [])

            if not documents:
                break

            all_results.extend(documents)
            offset += batch_size

        # Process all documents to get unique document names and their counts
        doc_info = {}
        for obj in all_results:
            doc_name = obj.get("document_name", "Unknown")  # Handle None case
            doc_type_obj = obj.get("doc_type", "Unknown")

            if doc_name not in doc_info:
                doc_info[doc_name] = {"type": doc_type_obj, "count": 0}
            doc_info[doc_name]["count"] += 1

        # Convert to list and sort
        all_documents = [{"name": name, "type": info["type"], "count": info["count"]} 
                         for name, info in doc_info.items()]
        all_documents.sort(key=lambda x: x["name"] or "")  # Handle None case in sorting

        # Pagination logic
        total_count = len(all_documents)
        total_pages = (total_count + per_page - 1) // per_page
        start_idx = (page - 1) * per_page
        end_idx = min(start_idx + per_page, total_count)
        documents = all_documents[start_idx:end_idx]

        logger.info(f"Found {total_count} documents, showing {len(documents)} on page {page}")
        logger.info(f"Document types found: {set(doc['type'] for doc in all_documents)}")
        logger.info(f"Documents found: {[doc['name'] for doc in all_documents]}")

        # Populate dropdown
        doc_types = get_unique_document_types()
        logger.info(f"Available document types: {doc_types}")

        # Add max and min functions to template context
        template_context = {
            'documents': documents,
            'current_page': page,
            'total_pages': total_pages,
            'per_page': per_page,
            'query': '',
            'doc_type': doc_type,
            'doc_types': doc_types,
            'all_documents': all_documents,  # Add all documents for dropdown
            'max': max,
            'min': min
        }

        return render_template('search.html', **template_context)
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return render_template(
            'search.html',
            documents=[],
            error=str(e),
            current_page=1,
            total_pages=1,
            per_page=10,
            query='',
            doc_type='',
            doc_types=[],
            all_documents=[],
            max=max,
            min=min
        )

# Initialize Weaviate client with cloud credentials
try:
    logger.info("Initializing Weaviate client...")
    # Use local Weaviate instance
    weaviate_url = "http://localhost:8080"
    
    # Initialize the client without authentication for local instance
    client = weaviate.Client(
        url=weaviate_url
    )
    
    # Test the connection
    client.schema.get()
    logger.info("Weaviate client initialized and connection tested successfully")
except Exception as e:
    logger.error(f"Failed to initialize Weaviate client: {str(e)}")
    client = None  # Set client to None instead of raising the error

# Gemini API configuration
GEMINI_API_KEY = os.getenv('GOOGLE_API_KEY')
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

def generate_gemini_response(prompt):
    """Generate response using Gemini API via REST"""
    try:
        headers = { 
            "Content-Type": "application/json" 
        }
        params = {
            "key": GEMINI_API_KEY
        }
        data = {
            "contents": [{
                "parts": [{
                    "text": prompt
                }]
            }]
        }
        
        logger.info("Making request to Gemini API...")
        response = requests.post(
            GEMINI_API_URL,
            headers=headers,
            params=params,
            json=data
        )
        
        if response.status_code == 200:
            result = response.json()
            if 'candidates' in result and len(result['candidates']) > 0:
                return result['candidates'][0]['content']['parts'][0]['text']
            else:
                logger.error("No valid response from Gemini API")
                logger.error(f"Response: {response.text}")
                return None
        else:
            logger.error(f"Gemini API request failed with status {response.status_code}: {response.text}")
            return None
            
    except Exception as e:
        logger.error(f"Error calling Gemini API: {str(e)}")
        return None

def summarize_with_gemini(text, query):
    """Summarize text using Gemini API with context from the query"""
    try:
        prompt = f"""
        Based on the following text and the user's query "{query}", provide a concise and relevant summary.
        Focus on information that directly relates to the query.
        
        Text:
        {text}
        
        Summary:
        """
        return generate_gemini_response(prompt) or text
    except Exception as e:
        logger.error(f"Error in Gemini summarization: {str(e)}")
        return text

def clean_text(text):
    """Clean text by removing formatting characters and extra spaces"""
    # Remove multiple dots and spaces
    text = ' '.join(text.split())
    # Remove sequences of dots
    text = text.replace('................................', ' ')
    text = text.replace('........................', ' ')
    text = text.replace('................', ' ')
    text = text.replace('........', ' ')
    text = text.replace('....', ' ')
    text = text.replace('..', ' ')
    # Remove extra spaces
    text = ' '.join(text.split())
    return text

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file, including form field values"""
    try:
        logger.info(f"Extracting text from PDF: {pdf_path}")
        reader = PdfReader(pdf_path)
        
        # Check if PDF is encrypted
        if reader.is_encrypted:
            try:
                # Try to decrypt with empty password first
                reader.decrypt('')
            except Exception as e:
                logger.error(f"PDF is encrypted and cannot be decrypted: {str(e)}")
                return None
        
        text = ""
        
        # First try to get form field values
        try:
            if reader.get_fields():
                logger.info("PDF contains form fields, extracting values...")
                fields = reader.get_fields()
                for field_name, field_value in fields.items():
                    if field_value and str(field_value).strip():
                        # Get the clean field name (either T or TU attribute)
                        field_obj = field_value
                        clean_name = field_obj.get('/T') or field_obj.get('/TU') or field_name
                        # Remove any PDF-specific prefixes or suffixes
                        clean_name = clean_name.replace('/T', '').replace('/TU', '').strip()
                        # Get the field value
                        field_value = field_obj.get('/V', '')
                        if field_value:
                            text += f"{clean_name}: {field_value}\n"
                logger.info(f"Extracted {len(fields)} form field values")
        except Exception as e:
            logger.warning(f"Could not extract form fields: {str(e)}")
        
        # Then extract regular text from pages
        for i, page in enumerate(reader.pages):
            try:
                logger.debug(f"Processing page {i+1}")
                page_text = page.extract_text()
                if page_text:
                    # Clean the extracted text
                    page_text = clean_text(page_text)
                    text += page_text + "\n"
            except Exception as e:
                logger.warning(f"Error processing page {i+1}: {str(e)}")
                continue
        
        if not text.strip():
            logger.warning("No text could be extracted from the PDF")
            return None
            
        logger.info(f"Successfully extracted text from {len(reader.pages)} pages")
        logger.debug(f"Extracted text: {text[:500]}...")
        return text
    except Exception as e:
        logger.error(f"Error reading PDF: {str(e)}")
        return None

def extract_text_from_file(file_path):
    """Extract text from a file (PDF or DOCX)"""
    try:
        logger.info(f"Extracting text from file: {file_path}")
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            text = extract_text_from_pdf(file_path)
            if not text:
                logger.error("Failed to extract text from PDF")
                return None
        elif ext == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                if not text.strip():
                    logger.warning("No text found in DOCX file")
                    return None
            except Exception as e:
                logger.error(f"Error reading DOCX file: {str(e)}")
                return None
        elif ext == 'doc':
            logger.error("Legacy .doc files are not supported. Please upload a .docx file.")
            return None
        else:
            logger.error(f"Unsupported file type: {file_path}")
            return None
        return clean_text(text)
    except Exception as e:
        logger.error(f"Error reading file: {str(e)}")
        return None

def create_chunks(text, chunk_size=5000):
    """Split text into chunks of approximately equal size"""
    try:
        logger.info(f"Creating chunks with size {chunk_size}") 
        chunks = textwrap.wrap(text, chunk_size, break_long_words=False, break_on_hyphens=False)
        logger.info(f"Created {len(chunks)} chunks")
        return chunks
    except Exception as e:
        logger.error(f"Error creating chunks: {str(e)}")
        raise

def create_schema():
    """Create the schema for the Document class in Weaviate."""
    try:
        if client is None:
            logger.error("Cannot create schema: Weaviate client is not initialized")
            return False
            
        # Check if Document class exists
        if not client.schema.exists("Document"):
            logger.info("Document class does not exist, creating schema...")
            document_schema = {
                "class": "Document",
                "description": "A class to store documents with metadata",
                "vectorizer": "text2vec-openai",
                "properties": [
                    {
                        "name": "document_name",
                        "dataType": ["string"],
                        "description": "The name of the document"
                    },
                    {
                        "name": "content",
                        "dataType": ["text"],
                        "description": "The content of the document"
                    },
                    {
                        "name": "doc_type",
                        "dataType": ["string"],
                        "description": "The type of document (Legal templates, Legal cases, General, FAQ)"
                    },
                    {
                        "name": "timestamp",
                        "dataType": ["string"],
                        "description": "When the document was uploaded"
                    }
                ]
            }
            
            client.schema.create_class(document_schema)
            logger.info("Created Document class schema")
        else:
            logger.info("Document class already exists")
            
        # Verify schema
        schema = client.schema.get("Document")
        logger.info(f"Current Document schema: {json.dumps(schema, indent=2)}")
        
        return True
    except Exception as e:
        logger.error(f"Error creating schema: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return False

def extract_court_case_metadata(text, doc_type):
    """Extract metadata from court case documents using Gemini AI."""
    try:
        if doc_type.lower() not in ['court case', 'legal case', 'legal cases']:
            return None

        prompt = f"""
        Extract the following metadata from this court case document. If a field is not found, use "Not specified".
        Return the data in JSON format.

        Required fields:
        - case_title: The full title of the case (e.g., "Matter of Cuppek v. DiNapoli")
        - citation: The official citation
        - court: The court name
        - jurisdiction: The jurisdiction
        - decision_date: The date of the decision
        - calendar_date: The calendar date
        - docket_number: The docket number
        - parties: The named parties involved
        - issue: The main issue of the case
        - outcome: The final outcome
        - judges: List of judges
        - authoring_judge: The judge who authored the opinion
        - petitioner_attorney: The petitioner's attorney details
        - respondent_attorney: The respondent's attorney details
        - publisher: The publisher
        - statutes_cited: List of statutes cited
        - prior_history: Prior case history
        - keywords: List of relevant keywords
        - document_status: Status of the document
        - source: Source of the document

        Document Content:
        {text}

        Return the data in this exact JSON format:
        {{
            "case_title": "string",
            "citation": "string",
            "court": "string",
            "jurisdiction": "string",
            "decision_date": "string",
            "calendar_date": "string",
            "docket_number": "string",
            "parties": "string",
            "issue": "string",
            "outcome": "string",
            "judges": ["string"],
            "authoring_judge": "string",
            "petitioner_attorney": "string",
            "respondent_attorney": "string",
            "publisher": "string",
            "statutes_cited": ["string"],
            "prior_history": "string",
            "keywords": ["string"],
            "document_status": "string",
            "source": "string"
        }}
        """

        response = generate_gemini_response(prompt)
        if not response:
            logger.error("Failed to extract metadata from court case")
            return None

        # Extract JSON from the response
        import re
        import json
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if not json_match:
            logger.error("Could not find JSON in metadata response")
            return None

        metadata = json.loads(json_match.group(0))
        logger.info(f"Extracted metadata for court case: {metadata.get('case_title', 'Unknown')}")
        return metadata

    except Exception as e:
        logger.error(f"Error extracting court case metadata: {str(e)}")
        return None

def process_and_upload_file(file_path, doc_type, chunk_size=5000):
    """Process a file (PDF or TXT) and upload its chunks to Weaviate"""
    try:
        logger.info(f"Starting file processing: {file_path}")
        logger.info(f"Document type: {doc_type}")
        logger.info(f"Chunk size: {chunk_size}")
        
        # Extract text from file
        logger.info("Extracting text from file...")
        text = extract_text_from_file(file_path)
        if not text:
            logger.error("Failed to extract text from file")
            return False
        logger.info(f"Successfully extracted text. Text length: {len(text)} characters")

        # Extract metadata for court cases
        metadata = None
        if doc_type.lower() in ['court case', 'legal case', 'legal cases']:
            logger.info("Extracting metadata for legal case document")
            metadata = extract_court_case_metadata(text, doc_type)
            if metadata:
                logger.info(f"Successfully extracted metadata: {json.dumps(metadata, indent=2)}")
            else:
                logger.warning("Failed to extract metadata from court case")
                return False

        # Create chunks
        logger.info("Creating chunks from extracted text...")
        chunks = create_chunks(text, chunk_size)
        logger.info(f"Created {len(chunks)} chunks")
        
        filename = os.path.basename(file_path)
        
        # Create schema for all documents
        logger.info("Creating/verifying Weaviate schema...")
        create_schema()
        
        # Check if document already exists
        existing_docs = (
            client.query
            .get("Document", ["document_name", "doc_type"])
            .with_where({
                "path": ["document_name"],
                "operator": "Equal",
                "valueString": filename
            })
            .do()
        )
        
        if existing_docs and "data" in existing_docs and "Get" in existing_docs["data"] and "Document" in existing_docs["data"]["Get"] and existing_docs["data"]["Get"]["Document"]:
            existing_doc = existing_docs["data"]["Get"]["Document"][0]
            existing_type = existing_doc.get('doc_type', 'Unknown')
            logger.warning(f"Document {filename} already exists with type {existing_type}. Deleting existing document...")
            
            # Delete existing document
            delete_document(filename)
            logger.info(f"Deleted existing document {filename}")
        
        logger.info(f"Uploading {len(chunks)} chunks to Weaviate collection: Document")
        for i, chunk in enumerate(chunks):
            try:
                logger.debug(f"Processing chunk {i+1}/{len(chunks)}")
                # Normalize document type
                normalized_doc_type = doc_type.strip()
                if normalized_doc_type.lower() == "legal templates":
                    normalized_doc_type = "Legal templates"
                elif normalized_doc_type.lower() == "legal cases":
                    normalized_doc_type = "Legal cases"
                elif normalized_doc_type.lower() == "texas constitution and statutes":
                    normalized_doc_type = "Texas Constitution and Statutes"
                
                logger.info(f"Using normalized document type: {normalized_doc_type}")
                
                data_object = {
                    "document_name": filename,
                    "chunk_index": i,
                    "content": chunk,
                    "doc_type": normalized_doc_type,
                    "timestamp": datetime.now().isoformat()
                }

                # Add metadata for court cases
                if metadata:
                    logger.info(f"Adding metadata to chunk {i+1}")
                    data_object.update(metadata)
                    logger.info(f"Data object with metadata for chunk {i+1}: {json.dumps(data_object, indent=2)}")

                logger.info(f"Uploading chunk {i+1} to Weaviate...")
                result = client.data_object.create(
                    class_name="Document",
                    data_object=data_object
                )
                logger.info(f"Successfully uploaded chunk {i+1} with type {normalized_doc_type}. Object ID: {result}")
                
                # Verify the upload
                verify_result = (
                    client.query
                    .get("Document", ["_id"])
                    .with_where({
                        "path": ["_id"],
                        "operator": "Equal",
                        "valueString": result
                    })
                    .do()
                )
                if verify_result and "data" in verify_result and "Get" in verify_result["data"] and "Document" in verify_result["data"]["Get"]:
                    uploaded_doc = verify_result["data"]["Get"]["Document"][0]
                    logger.info(f"Verified upload - Document: {uploaded_doc['document_name']}, Type: {uploaded_doc.get('doc_type', 'Unknown')}")
                else:
                    logger.warning(f"Could not verify upload for chunk {i+1}")
            except Exception as e:
                logger.error(f"Error uploading chunk {i}: {str(e)}")
                logger.error(f"Stack trace: {traceback.format_exc()}")
                return False
        
        logger.info("File processing and upload completed successfully")
        return True
    except Exception as e:
        logger.error(f"Error in process_and_upload_file: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return False

def format_text_with_bold(text):
    """Remove any markdown formatting from text"""
    try:
        # Remove ** and * formatting
        text = text.replace('**', '').replace('*', '')
        return text
    except Exception as e:
        logger.error(f"Error formatting text: {str(e)}")
        return text

def perform_search(query, document_name=None):
    """Perform a text-based search on the collection using bm25 and generate a focused answer using Gemini API."""
    import traceback
    GENERAL_QUERIES = [
        "what is this document about",
        "summarize",
        "what are main topics of this document",
        "summary",
        "main topics",
        "overview",
        "describe this document",
        "what does this document contain"
    ]
    try:
        logger.info(f"Performing bm25 search for query: '{query}' in document: {document_name}")
        query_clean = clean_text(query).lower().strip(" ?.")
        is_general = any(gq in query_clean for gq in GENERAL_QUERIES)
        # Clean the search query
        query = clean_text(query)
        try:
            logger.debug(f"Querying Document with query: '{query}', document_name: '{document_name}'")
            
            # Create the query builder
            query_builder = client.query.get("Document", ["document_name", "chunk_index", "content", "doc_type"])
            
            # Add BM25 search
            query_builder = query_builder.with_bm25(
                    query=query,
                    properties=["content"]
                )
            
            # Add where filter if provided
            if document_name:
                if isinstance(document_name, dict):
                    # If it's already a filter object, use it directly
                    query_builder = query_builder.with_where(document_name)
                elif isinstance(document_name, str):
                    # Create a simple equality filter for document name
                    query_builder = query_builder.with_where({
                        "path": ["document_name"],
                        "operator": "Equal",
                        "valueString": document_name
                    })
            
            # Execute the query
            result = query_builder.with_limit(5).do()
            
            logger.debug(f"Raw Weaviate response: {json.dumps(result, indent=2)}")
            results = []
            if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
                results = result["data"]["Get"]["Document"]
                logger.info(f"Initial results count: {len(results)}")
                # Filter by document_name if specified
                if document_name and isinstance(document_name, str):
                    results = [r for r in results if r["document_name"] == document_name]
                    logger.info(f"Results after filtering by document_name '{document_name}': {len(results)}")
                for r in results:
                    r["content"] = clean_text(r["content"])
            # If general query or no results, fetch all chunks for the document
            if is_general or not results:
                logger.info("General query or no results, fetching all chunks for document.")
                all_chunks_result = (
                    client.query
                    .get("Document", ["document_name", "chunk_index", "content", "doc_type"])
                    .with_where(document_name if isinstance(document_name, dict) else None)
                    .with_limit(100)
                    .do()
                )
                all_chunks = []
                if all_chunks_result and "data" in all_chunks_result and "Get" in all_chunks_result["data"] and "Document" in all_chunks_result["data"]["Get"]:
                    all_chunks = all_chunks_result["data"]["Get"]["Document"]
                logger.info(f"Fetched {len(all_chunks)} chunks for Gemini summarization.")
                if all_chunks:
                    combined_content = "\n\n".join([c["content"] for c in all_chunks])
                    prompt = f"""
                    Please provide a summary or main topics for the following document content. If the user query is:\n'{query}', answer accordingly.\n\nContent:\n{combined_content}\n\nSummary/Main Topics/Overview:
                    """
                    focused_answer = generate_gemini_response(prompt)
                    if focused_answer:
                        # Parse and format the answer
                        soup = BeautifulSoup(focused_answer, 'html.parser')
                        text = soup.get_text()
                        # Replace markdown-style formatting
                        text = text.replace("• **", "\n- ").replace(":**", ":").replace("**", "")
                        # Add proper line breaks
                        text = text.replace("\n\n", "<br><br>").replace("\n", "<br>")
                        logger.info("Returning Gemini summary for general query or no results.")
                        return [{
                            "document_name": document_name or "Answer",
                            "content": text,
                            "summary": text
                        }]
            # If results exist and not a general query, use Gemini on those chunks
            if results:
                combined_content = "\n\n".join([r["content"] for r in results])
                # Group results by document type
                doc_types = set(r.get("doc_type", "Unknown") for r in results)
                doc_type_info = f"Found in document types: {', '.join(doc_types)}" if doc_types else ""
                
                prompt = f"""
                Based on the following information, provide a clear and well-structured answer to the question: "{query}"
                {doc_type_info}
                Your answer should:
                1. Start with a brief introduction (1-2 sentences)
                2. Use bullet points (•) for key information when there are multiple points
                3. Include 3-5 main points
                4. End with a brief conclusion if relevant
                5. Use proper line breaks between sections
                Format your response like this:
                [Introduction]
                • Point 1
                • Point 2
                • Point 3
                [Conclusion if needed]
                Information:
                {combined_content}
                Structured Answer:
                """
                focused_answer = generate_gemini_response(prompt)
                if focused_answer:
                    # Parse and format the answer
                    soup = BeautifulSoup(focused_answer, 'html.parser')
                    text = soup.get_text()
                    # Replace markdown-style formatting
                    text = text.replace("• **", "\n- ").replace(":**", ":").replace("**", "")
                    # Add proper line breaks
                    text = text.replace("\n\n", "<br><br>").replace("\n", "<br>")
                    logger.info("Returning focused Gemini answer.")
                    return [{
                        "document_name": "Answer",
                        "content": text,
                        "summary": text
                    }]
                logger.info(f"Found {len(results)} matching results (after all filtering)")
                return results
            else:
                logger.warning(f"No results found in the expected format. Raw response: {json.dumps(result, indent=2)}")
        except Exception as e:
            logger.error(f"Error searching collection: {str(e)}\n{traceback.format_exc()}")
            return []
        logger.warning("No results found in the expected format after query block.")
        return []
    except Exception as e:
        logger.error(f"Error performing search: {str(e)}\n{traceback.format_exc()}")
        return []

def delete_collection(collection_name):
    """Delete a collection and all its objects from Weaviate"""
    try:
        logger.info(f"Deleting collection: {collection_name}")
        
        # First delete all objects in the collection
        try:
            result = (
                client.query
                .get(collection_name, ["_additional {id}"])
                .do()
            )
            if result and "data" in result and "Get" in result["data"] and collection_name in result["data"]["Get"]:
                objects = result["data"]["Get"][collection_name]
                for obj in objects:
                    obj_id = obj["_additional"]["id"]
                    client.data_object.delete(
                        class_name=collection_name,
                        uuid=obj_id
                    )
                logger.info(f"Deleted {len(objects)} objects from collection {collection_name}")
        except Exception as e:
            logger.error(f"Error deleting objects from collection {collection_name}: {str(e)}")
            return False
        
        # Then delete the collection schema
        try:
            client.schema.delete_class(collection_name)
            logger.info(f"Successfully deleted collection schema: {collection_name}")
            return True
        except Exception as e:
            logger.error(f"Error deleting collection schema {collection_name}: {str(e)}")
            return False
            
    except Exception as e:
        logger.error(f"Error in delete_collection: {str(e)}")
        return False

def get_collections():
    """Get all available collections from Weaviate."""
    try:
        # Get the schema from Weaviate
        schema = client.schema.get()
        
        # Extract class names (collections) from the schema
        collections = [class_info["class"] for class_info in schema.get("classes", [])]
        
        # Log the number of collections found
        logger.info(f"Found {len(collections)} collections: {collections}")
        
        return collections
    except Exception as e:
        logger.error(f"Error getting collections: {str(e)}")
        raise

@app.route('/upload', methods=['POST'])
@login_required
def upload():
    logger.debug("Starting upload route")
    logger.debug(f"Request method: {request.method}")
    logger.debug(f"Request form data: {request.form}")
    logger.debug(f"Request files: {request.files}")
    
    if client is None:
        logger.error("Weaviate client is not initialized")
        return redirect(url_for('manage_collections', 
                            message="Weaviate connection is not available. Please check your configuration.", 
                            message_type="error"))
                            
    if 'file' not in request.files:
        logger.warning("No file part in request")
        return redirect(url_for('manage_collections', 
                            message="No file selected", 
                            message_type="error"))
    
    file = request.files['file']
    doc_type = request.form.get('docType')
    
    logger.debug(f"Received file: {file.filename}")
    logger.debug(f"Received document type: {doc_type}")
    
    if file.filename == '':
        logger.warning("No file selected")
        return redirect(url_for('manage_collections', 
                            message="No file selected", 
                            message_type="error"))
    
    if not doc_type:
        logger.warning("No document type selected")
        return redirect(url_for('manage_collections', 
                            message="Please select a document type", 
                            message_type="error"))
    
    logger.info(f"Starting upload process for file: {file.filename}")
    logger.info(f"Raw document type from form: {doc_type}")
    logger.info(f"Document type after strip(): {doc_type.strip()}")
    logger.info(f"Document type lower(): {doc_type.lower()}")
    logger.info(f"Document type title(): {doc_type.title()}")
    
    # Normalize document type
    normalized_doc_type = doc_type.strip()
    if normalized_doc_type.lower() == "legal templates":
        normalized_doc_type = "Legal templates"
    elif normalized_doc_type.lower() == "legal cases":
        normalized_doc_type = "Legal cases"
    elif normalized_doc_type.lower() == "texas constitution and statutes":
        normalized_doc_type = "Texas Constitution and Statutes"
    
    logger.info(f"Normalized document type: {normalized_doc_type}")
    
    if file and (file.filename.lower().endswith('.pdf') or file.filename.lower().endswith('.txt')):
        try:
            # Save the file temporarily
            filename = secure_filename(file.filename)
            temp_path = os.path.join('temp', filename)
            os.makedirs('temp', exist_ok=True)
            logger.info(f"Saving file temporarily to: {temp_path}")
            file.save(temp_path)
            logger.info(f"File saved successfully to temporary location")
            
            # Process and upload the file
            logger.info(f"Starting file processing and upload to Weaviate with type: {normalized_doc_type}")
            if process_and_upload_file(temp_path, normalized_doc_type):
                logger.info(f"Successfully processed and uploaded {filename} with type {normalized_doc_type}")
               
                os.remove(temp_path)
                logger.info("Temporary file removed")
                return redirect(url_for('manage_collections', 
                                    message=f"Successfully uploaded and processed {filename}", 
                                    message_type="success"))
            else:
                logger.error(f"Failed to process {filename}")
                # Clean up the temporary file
                os.remove(temp_path)
                logger.info("Temporary file removed after processing failure")
                return redirect(url_for('manage_collections', 
                                    message=f"Failed to process {filename}", 
                                    message_type="error"))
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            logger.error(f"Stack trace: {traceback.format_exc()}")
            return redirect(url_for('manage_collections', 
                                message=f"Error processing file: {str(e)}", 
                                message_type="error"))
    else:
        logger.warning(f"Invalid file type: {file.filename}")
        return redirect(url_for('manage_collections', 
                            message="Please upload a PDF or TXT file", 
                            message_type="error"))

@app.route('/search', methods=['POST'])
@login_required
def search():
    query = request.form.get('query', '')
    document_name = request.form.get('document', '')
    doc_type = request.form.get('doc_type', '')
    logger.info(f"Handling search request for query: '{query}' in document: {document_name} with type: {doc_type}")
    
    if not query:
        return jsonify({
            'error': 'Please enter a search query'
        }), 400
    
    if client is None:
        logger.error("Weaviate client is not initialized")
        return jsonify({
            'error': 'Weaviate connection is not available. Please check your configuration.'
        }), 503
    
    try:
        # Prepare filter conditions
        filter_conditions = []
        
        # Add document name filter if specified
        if document_name:
            filter_conditions.append({
                "path": ["document_name"],
                "operator": "Equal",
                "valueString": document_name
            })
        
        # Add document type filter if specified
        if doc_type:
            filter_conditions.append({
                "path": ["doc_type"],
                "operator": "Equal",
                "valueString": doc_type
            })
        
        # Combine all filters with AND if any filters exist
        where_filter = {
            "operator": "And",
            "operands": filter_conditions
        } if filter_conditions else None
        
        # Perform search with filters
        results = perform_search(query, where_filter)
        if not results:
            return jsonify({
                'results': [],
                'message': f"No results found for '{query}'"
            })
        
        # Convert results to a format suitable for JSON serialization
        serialized_results = []
        for result in results:
            serialized_result = {
                'document_name': result.get('document_name', ''),
                'chunk_index': result.get('chunk_index'),
                'content': result.get('content', ''),
                'summary': result.get('summary', ''),
                'doc_type': result.get('doc_type', '')  # Add document type to results
            }
            serialized_results.append(serialized_result)
        
        return jsonify({
            'results': serialized_results,
            'message': f"Found {len(results)} results for '{query}'"
        })
        
    except Exception as e:
        logger.error(f"Error in search route: {str(e)}")
        return jsonify({
            'error': f"Search error: {str(e)}"
        }), 500

@app.route('/delete', methods=['POST'])
@login_required
def delete():
    collection_name = request.form.get('collection', '')
    if not collection_name:
        return jsonify({
            'error': 'No collection specified'
        }), 400
    
    try:
        if delete_collection(collection_name):
            return jsonify({
                'message': f'Successfully deleted collection {collection_name}'
            })
        else:
            return jsonify({
                'error': f'Failed to delete collection {collection_name}'
            }), 500
    except Exception as e:
        logger.error(f"Error in delete route: {str(e)}")
        return jsonify({
            'error': f"Delete error: {str(e)}"
        }), 500

@app.route('/manage')
@login_required
def manage_collections():
    try:
        if client is None:
            logger.error("Weaviate client is not initialized")
            return render_template('manage_collections.html', 
                                documents=[], 
                                error="Weaviate connection is not available. Please check your configuration.",
                                current_page=1,
                                total_pages=1,
                                per_page=10,
                                doc_type='',
                                doc_types=[])
        
        # Get query parameters
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 10))
        doc_type = request.args.get('doc_type', '')
        
        logger.info(f"Fetching documents with type filter: {doc_type}")
        
        # Prepare type filter
        filter_clause = None
        if doc_type:
            doc_type = doc_type.replace('%20', ' ')
            logger.info(f"Searching for documents with type: {doc_type}")

            filter_clause = {
                        "path": ["doc_type"],
                        "operator": "Equal",
                "valueString": doc_type
            }

        # Fetch all documents using offset-based pagination
        offset = 0
        batch_size = 2000
        all_results = []

        while True:
            query = (
                client.query
                .get("Document", ["document_name", "doc_type"])
                .with_additional("id")
                .with_limit(batch_size)
                .with_offset(offset)
            )

            if filter_clause:
                query = query.with_where(filter_clause)

            result = query.do()
            documents = result.get("data", {}).get("Get", {}).get("Document", [])

            if not documents:
                break

            all_results.extend(documents)
            offset += batch_size
        
        # Process all documents to get unique document names and their counts
        doc_info = {}
        for obj in all_results:
            doc_name = obj.get("document_name", "Unknown")  # Handle None case
            doc_type_obj = obj.get("doc_type", "Unknown")

            if doc_name not in doc_info:
                doc_info[doc_name] = {"type": doc_type_obj, "count": 0}
            doc_info[doc_name]["count"] += 1
            
        # Convert to list and sort
        all_documents = [{"name": name, "type": info["type"], "count": info["count"]} 
                        for name, info in doc_info.items()]
        all_documents.sort(key=lambda x: x["name"] or "")  # Handle None case in sorting

        # Pagination logic
        total_count = len(all_documents)
        total_pages = (total_count + per_page - 1) // per_page
        start_idx = (page - 1) * per_page
        end_idx = min(start_idx + per_page, total_count)
        documents = all_documents[start_idx:end_idx]

        logger.info(f"Found {total_count} documents, showing {len(documents)} on page {page}")
        logger.info(f"Document types found: {set(doc['type'] for doc in all_documents)}")
        logger.info(f"Documents found: {[doc['name'] for doc in all_documents]}")

        # Populate dropdown
        doc_types = get_unique_document_types()
        logger.info(f"Available document types: {doc_types}")

        # Add max and min functions to template context
        template_context = {
            'documents': documents,
            'current_page': page,
            'total_pages': total_pages,
            'per_page': per_page,
            'doc_type': doc_type,
            'doc_types': doc_types,
            'max': max,
            'min': min
        }

        return render_template('manage_collections.html', **template_context)
    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return render_template(
            'manage_collections.html',
            documents=[],
            error=str(e),
            current_page=1,
            total_pages=1,
            per_page=10,
            doc_type='',
            doc_types=[],
            max=max,
            min=min
        )


@app.route('/delete_document/<document_name>', methods=['DELETE'])
@login_required
def delete_document(document_name):
    """Delete all objects for a specific document from Weaviate."""
    try:
        # First get all objects for this document
        result = (
            client.query
            .get("Document", ["_additional {id}"])
            .with_where({
                "path": ["document_name"],
                "operator": "Equal",
                "valueString": document_name
            })
            .do()
        )
        
        if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
            objects = result["data"]["Get"]["Document"]
            for obj in objects:
                obj_id = obj["_additional"]["id"]
                client.data_object.delete(
                    class_name="Document",
                    uuid=obj_id
                )
            logger.info(f"Deleted {len(objects)} objects for document {document_name}")
            return '', 204  # No content response for successful deletion
        else:
            return jsonify({'error': f'Document {document_name} not found'}), 404
            
    except Exception as e:
        logger.error(f"Error deleting document {document_name}: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/delete_object/<object_id>', methods=['DELETE'])
@login_required
def delete_object(object_id):
    """Delete a specific object from Weaviate."""
    try:
        logger.info(f"Deleting object with ID: {object_id}")
        client.data_object.delete(
            class_name="Document",
            uuid=object_id
        )
        logger.info(f"Successfully deleted object {object_id}")
        return '', 204  # No content response for successful deletion
    except Exception as e:
        logger.error(f"Error deleting object {object_id}: {str(e)}")
        return jsonify({'error': str(e)}), 500

def search_vector_collection(query, limit=5):
    """Search the vector collection for relevant content."""
    try:
        logger.info(f"Performing vector search for query: '{query}'")
        
        # Clean and prepare the search query
        search_terms = query.lower().split()
        logger.info(f"Search terms: {search_terms}")
        
        # Search in the Document with multiple search strategies
        results = []
        
        # Try exact match first
        result = (
            client.query
            .get("Document", ["document_name", "content"])
            .with_bm25(
                query=query,
                properties=["content", "document_name"]
            )
            .with_limit(limit)
            .do()
        )
        
        if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
            results.extend(result["data"]["Get"]["Document"])
            logger.info(f"Found {len(results)} results from exact match")
        
        # If not enough results, try vector search
        if len(results) < 2:
            result = (
                client.query
                .get("Document", ["document_name", "content"])
                .with_near_text({"concepts": [query]})
                .with_limit(limit)
                .do()
            )
            
            if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
                vector_results = result["data"]["Get"]["Document"]
                # Add only unique results
                for vr in vector_results:
                    if not any(r["document_name"] == vr["document_name"] for r in results):
                        results.append(vr)
                logger.info(f"Added {len(vector_results)} results from vector search")
        
        # Log the found documents
        if results:
            logger.info("Found documents:")
            for r in results:
                logger.info(f"- {r['document_name']}")
        else:
            logger.warning("No documents found in the collection")
            
        return results
    except Exception as e:
        logger.error(f"Error searching vector collection: {str(e)}")
        return []

def generate_documentation_with_ai(title, requirements, doc_type):
    """Generate legal documentation using Gemini API."""
    try:
        # Search vector collection for relevant content
        vector_results = search_vector_collection_by_type(requirements, doc_type)
        # Log the search results
        logger.info(f"Found {len(vector_results)} relevant documents")
        if vector_results:
            logger.info("Found documents:")
            for r in vector_results:
                logger.info(f"- {r['document_name']} (Type: {r.get('doc_type', 'Unknown')})")
        
        # Check if we have enough relevant data
        if not vector_results:
            logger.warning("No relevant documents found in the vector collection")
            return """
            <div class="alert alert-warning">
                <h3>No Relevant Data Found</h3>
                <p>We could not find any relevant data in our vector collection for your request. 
                Please try:</p>
                <ul>
                    <li>Refining your search requirements</li>
                    <li>Adding more specific details about the document you need</li>
                    <li>Uploading relevant documents to the collection first</li>
                </ul>
                <p>Note: We require relevant data from our verified vector collection 
                to ensure accuracy and reliability of the generated document.</p>
            </div>
            """
        
        # Combine all relevant content
        vector_knowledge = "\n\n".join([result["content"] for result in vector_results])
        
        # Use Gemini API to generate the document
        prompt = f"""
You are an expert legal document generator.

When generating any section that contains structured data (such as parties, property details, or any form-like information), use HTML <table>, <tr>, and <td> tags. Each label and value should be in its own cell. Do NOT use manual spaces or multiple <span> tags for alignment.

For example:
<table>
  <tr><td>Landlord Name:</td><td>Swathi</td></tr>
  <tr><td>Tenant Name:</td><td>Srikanth</td></tr>
  <tr><td>Mailing Address:</td><td>213123, Frisco, Texas, USA</td></tr>
</table>

Use headings for section titles. Format all structured/form sections using HTML tables.

Document Type: {doc_type}
Title: '{title}'

Requirements:
{requirements}

Relevant Knowledge:
{vector_knowledge}

The legal template points to keep exactly as-is:

1. Be well-structured and professionally formatted
2. Include all necessary legal clauses and sections
3. Be suitable for professional use
4. Follow standard legal document conventions
5. Include proper headings and sections
6. Use clear and precise language

Format the entire generated document using HTML tags for proper display.

Generate the document now, updating only with the given dynamic data and instructions but preserving the template points verbatim.
"""
        
        document_content = generate_gemini_response(prompt)
        if not document_content:
            logger.error("Gemini API did not return a document.")
            return "<div class='alert alert-danger'>Failed to generate document using Gemini AI.</div>"
            
        return document_content
    except Exception as e:
        logging.error(f"Error generating documentation: {str(e)}")
        return f"Error generating documentation: {str(e)}"

def store_document_in_weaviate(title, content, doc_type):
    """Store the generated document in Weaviate."""
    try:
        logger.info(f"Storing document in Weaviate: {title}")
        
        # Create document object
        document_object = {
            "title": title,
            "content": content,
            "type": doc_type,
            "timestamp": datetime.now().isoformat()
        }
        
        # Store in Weaviate
        client.data_object.create(
            class_name="Document",
            data_object=document_object
        )
        
        logger.info(f"Successfully stored document: {title}")
        return True
    except Exception as e:
        logger.error(f"Error storing document in Weaviate: {str(e)}")
        return False

def get_unique_document_types():
    """Get unique document types from Weaviate"""
    try:
        # Get all documents without filtering
        result = (
            client.query
            .get("Document", ["doc_type"])
            .with_additional("id")
            .do()
        )
        
        doc_types = set()
        if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
            for obj in result["data"]["Get"]["Document"]:
                if "doc_type" in obj and obj["doc_type"]:
                    # Normalize document type
                    doc_type = obj["doc_type"].strip()
                    if doc_type.lower() == "legal templates":
                        doc_type = "Legal templates"
                    elif doc_type.lower() == "legal cases":
                        doc_type = "Legal cases"
                    elif doc_type.lower() == "texas constitution and statutes":
                        doc_type = "Texas Constitution and Statutes"
                    doc_types.add(doc_type)
                    logger.info(f"Found document type: {doc_type}")
        
        # Add default document types if they don't exist
        default_types = {
            "Legal templates",
            "Legal cases",
            "General",
            "Texas Constitution and Statutes",
            "FAQ"
        }
        doc_types.update(default_types)
        
        logger.info(f"All document types: {sorted(list(doc_types))}")
        return sorted(list(doc_types))
    except Exception as e:
        logger.error(f"Error getting document types: {str(e)}")
        return []

@app.route('/document_creation')
@login_required
def document_creation():
    """Render the document creation dashboard."""
    try:
        logger.info("Starting document_creation route")
        
        # Get all documents with pagination
        offset = 0
        batch_size = 1000
        all_documents = []
        
        while True:
            result = (
                client.query
                .get("Document", ["document_name", "doc_type"])
                .with_additional("id")
                .with_where({
                    "path": ["doc_type"],
                    "operator": "Equal",
                    "valueString": "Legal templates"
                })
                .with_limit(batch_size)
                .with_offset(offset)
                .do()
            )
            
            if not result or "data" not in result or "Get" not in result["data"] or "Document" not in result["data"]["Get"]:
                break
                
            documents = result["data"]["Get"]["Document"]
            if not documents:
                break
                
            all_documents.extend(documents)
            offset += batch_size
        
        logger.info(f"Total legal template documents found: {len(all_documents)}")
        
        # Process documents to get unique legal templates
        legal_templates = {}
        
        for obj in all_documents:
            doc_name = obj.get("document_name", "Unknown")
            doc_type = obj.get("doc_type", "Unknown")
            
            logger.info(f"Processing document: {doc_name} with type: {doc_type}")
            
            if doc_type.strip().lower() == "legal templates":
                # Only add if we haven't seen this document name before
                if doc_name not in legal_templates:
                    legal_templates[doc_name] = {"name": doc_name, "type": doc_type}
                    logger.info(f"Added legal template: {doc_name}")
        
        # Convert dictionary to list and sort by name
        legal_templates_list = list(legal_templates.values())
        legal_templates_list.sort(key=lambda x: x["name"])
        
        logger.info(f"Final unique legal templates: {json.dumps(legal_templates_list, indent=2)}")
        
        return render_template('document_creation.html', 
                             legal_templates=legal_templates_list,
                             other_types=[])
    except Exception as e:
        logger.error(f"Error in document creation route: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return render_template('document_creation.html', 
                             legal_templates=[],
                             other_types=[],
                             error=str(e))

def search_vector_collection_by_type(query, doc_type, limit=5):
    """Search the vector collection for relevant content filtered by document type."""
    try:
        logger.info(f"Performing vector search for query: '{query}' with type: '{doc_type}'")
        
        # Clean and prepare the search query
        search_terms = query.lower().split()
        logger.info(f"Search terms: {search_terms}")
        
        # Search in the Document collection with type filter
        result = (
            client.query
            .get("Document", ["document_name", "content", "doc_type"])
            .with_bm25(
                query=query,
                properties=["content"]
            )
            .with_where({
                "operator": "And",
                "operands": [
                    {
                "path": ["doc_type"],
                        "operator": "Equal",
                        "valueString": "Legal templates"
                    },
                    {
                        "path": ["document_name"],
                "operator": "Equal",
                "valueString": doc_type
                    }
                ]
            })
            .with_limit(limit)
            .do()
        )
        
        if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
            return result["data"]["Get"]["Document"]
        return []
    except Exception as e:
        logger.error(f"Error searching vector collection: {str(e)}")
        return []

@app.route('/generate_document', methods=['POST'])
@login_required
def generate_document():
    """Generate documentation based on user requirements."""
    try:
        data = request.get_json()
        title = data.get('title')
        requirements = data.get('requirements')
        doc_type = data.get('docType')
        
        if not all([title, requirements, doc_type]):
            return jsonify({'error': 'Missing required fields'}), 400
        
        # Search vector collection for relevant content filtered by document type
        vector_results = search_vector_collection_by_type(requirements, doc_type)
        vector_sources = [{"document_name": result["document_name"]} for result in vector_results]
        
        # Generate documentation using AI
        document_content = generate_documentation_with_ai(title, requirements, doc_type)
        
        # Store in Weaviate
        store_document_in_weaviate(title, document_content, doc_type)
        
        # Return response with sources
        return jsonify({
            'document': document_content,
            'vector_sources': vector_sources,
            'external_sources': ["Gemini AI Model"]
        })
    except Exception as e:
        logging.error(f"Error in generate_document: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download_generated_document', methods=['POST'])
@login_required
def download_generated_document():
    """Generate and download the document as DOCX with improved formatting and invisible table borders."""
    try:
        data = request.get_json()
        content = data.get('content')
        title = data.get('title')
        
        if not content or not title:
            return jsonify({'error': 'Missing content or title'}), 400
        
        # Create a BytesIO buffer for the DOCX file
        buffer = io.BytesIO()
        
        # Create a new Document
        doc = Document()
        
        # Set document styles
        styles = doc.styles
        style = styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(10)
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph.space_after = Pt(20)
        
        # Convert HTML content to DOCX paragraphs and tables
        soup = BeautifulSoup(content, 'html.parser')
        
        def set_table_no_borders(table):
            table.style = 'Table Grid'
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    borders = OxmlElement('w:tcBorders')
                    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'none')
                        borders.append(border)
                    tcPr.append(borders)
        
        def add_kv_table(pairs):
            table = doc.add_table(rows=len(pairs), cols=2)
            set_table_no_borders(table)
            for i, (k, v) in enumerate(pairs):
                table.cell(i, 0).text = k
                table.cell(i, 1).text = v
            doc.add_paragraph()  # Add space after table
        
        for element in soup.find_all(['h2', 'h3', 'p', 'ul', 'li', 'div', 'table']):
            if element.name in ['h2', 'h3']:
                # Add heading
                p = doc.add_paragraph()
                p.style = 'Heading 1' if element.name == 'h2' else 'Heading 2'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(element.get_text().strip())
                p.space_after = Pt(12)
            elif element.name == 'p':
                # Add paragraph
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(element.get_text().strip())
            elif element.name == 'ul':
                # Start bullet list
                for li in element.find_all('li'):
                    p = doc.add_paragraph()
                    p.style = 'List Bullet'
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.add_run(li.get_text().strip())
            elif element.name == 'li':
                # Handle standalone list items
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(element.get_text().strip())
            elif element.name == 'table':
                # Convert HTML table to docx table
                rows = element.find_all('tr')
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                    set_table_no_borders(table)
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            table.cell(i, j).text = cell.get_text().strip()
                    doc.add_paragraph()
            elif element.name == 'div':
                # Try to detect key-value pairs in divs (e.g., <div><span>Label:</span><span>Value</span></div>)
                spans = element.find_all('span')
                if len(spans) >= 2 and len(spans) % 2 == 0:
                    pairs = []
                    for i in range(0, len(spans), 2):
                        k = spans[i].get_text().strip()
                        v = spans[i+1].get_text().strip()
                        pairs.append((k, v))
                    add_kv_table(pairs)
                else:
                    # Fallback: treat as paragraph
                    text = element.get_text().strip()
                    if text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p.add_run(text)
        
        # Save the document to the buffer
        doc.save(buffer)
        buffer.seek(0)
        
        # Create a response with the DOCX file
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{title.replace(' ', '_')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/debug_chunks/<document_name>')
@login_required
def debug_chunks(document_name):
    try:
        logger.info(f"Debug: Fetching all chunks for document_name: {document_name}")
        result = (
            client.query
            .get("Document", ["document_name", "chunk_index", "content"])
            .with_where({
                "path": ["document_name"],
                "operator": "Equal",
                "valueString": document_name
            })
            .with_limit(100)
            .do()
        )
        chunks = []
        if result and "data" in result and "Get" in result["data"] and "Document" in result["data"]["Get"]:
            chunks = result["data"]["Get"]["Document"]
        logger.info(f"Found {len(chunks)} chunks for document_name: {document_name}")
        return jsonify({"chunks": chunks, "count": len(chunks)})
    except Exception as e:
        logger.error(f"Error in debug_chunks: {str(e)}")
        return jsonify({"error": str(e)})

def search_court_cases(query, filters=None):
    try:
        if not client:
            logger.error("Weaviate client not initialized")
            return {"error": "Weaviate client not initialized"}, 500

        logger.info(f"Starting search with query: {query}")
        logger.info(f"Filters: {filters}")

        # Create base filter for document types
        doc_type_filter = {
            "operator": "Or",
            "operands": [
                {"path": ["doc_type"], "operator": "Equal", "valueText": "court case"},
                {"path": ["doc_type"], "operator": "Equal", "valueText": "legal case"},
                {"path": ["doc_type"], "operator": "Equal", "valueText": "legal cases"}
            ]
        }

        # Combine with additional filters if provided
        if filters:
            filter_conditions = [doc_type_filter]
            
            # Handle each filter based on its data type
            for field, value in filters.items():
                if value:  # Only add non-empty filters
                    logger.info(f"Adding filter for field: {field} with value: {value}")
                    if field in ["judges", "keywords", "statutes_cited"]:
                        # Array fields use ContainsAny operator
                        filter_conditions.append({
                            "path": [field],
                            "operator": "ContainsAny",
                            "valueText": [value] if isinstance(value, str) else value
                        })
                    elif field == "decision_date":
                        # Date field - use Equal for exact match
                        filter_conditions.append({
                            "path": [field],
                            "operator": "Equal",
                            "valueText": value
                        })
                    else:
                        # All other text fields use Equal for exact match
                        filter_conditions.append({
                            "path": [field],
                            "operator": "Equal",
                            "valueText": value
                        })
        
        # Combine all conditions with AND
            filter_clause = {
            "operator": "And",
            "operands": filter_conditions
        }
        else:
            filter_clause = doc_type_filter
        
        logger.info(f"Final filter clause: {json.dumps(filter_clause, indent=2)}")

        # Get metadata fields
        metadata_fields = [
            "document_name", "content", "case_title", "citation", "court",
            "jurisdiction", "decision_date", "docket_number", "parties",
            "judges", "authoring_judge", "keywords", "statutes_cited",
            "calendar_date", "petitioner_attorney", "respondent_attorney",
            "document_status", "source", "issue", "outcome", "prior_history"
        ]

        # Define searchable properties for BM25
        searchable_properties = [
            "content", "case_title", "citation", "court", 
            "jurisdiction", "parties", "judges", "authoring_judge",
            "keywords", "statutes_cited", "issue", "outcome", "prior_history"
        ]

        # Initialize results list
        results = []
        offset = 0
        batch_size = 1000

        while True:
            try:
                # BM25 search
                bm25_query = client.query.get(
                    "Document",
                    metadata_fields
                )

                # Only add BM25 if there's a query
                if query and query.strip():
                    bm25_query = bm25_query.with_bm25(
                    query=query,
                        properties=searchable_properties
                    )

                # Add filters
                bm25_query = bm25_query.with_where(
                    filter_clause
                ).with_limit(batch_size).with_offset(offset)

                logger.info(f"Executing BM25 search with offset {offset}")
                bm25_results = bm25_query.do()
                logger.info(f"BM25 results: {json.dumps(bm25_results, indent=2)}")

                # Vector search - only if we have a query
                try:
                    if query and query.strip():
                        # First try to get a vector representation
                        bm25_for_vector = client.query.get(
                            "Document", 
                            ["content", "_additional {vector}"]
                        ).with_bm25(
                            query=query,
                            properties=["content"]
                        ).with_limit(1).do()

                        if (bm25_for_vector and 
                            "data" in bm25_for_vector and 
                            "Get" in bm25_for_vector["data"] and 
                            "Document" in bm25_for_vector["data"]["Get"] and 
                            bm25_for_vector["data"]["Get"]["Document"]):
                            
                            vector = bm25_for_vector["data"]["Get"]["Document"][0]["_additional"]["vector"]
                            
                            vector_query = client.query.get(
                                "Document",
                                metadata_fields
                            ).with_near_vector({
                                "vector": vector
                            }).with_where(
                                filter_clause
                            ).with_limit(batch_size).with_offset(offset)

                            logger.info(f"Executing vector search with offset {offset}")
                            vector_results = vector_query.do()
                            logger.info(f"Vector results: {json.dumps(vector_results, indent=2)}")
                        else:
                            logger.warning("No results from BM25 query for vector search, skipping vector search")
                            vector_results = None
                    else:
                        logger.info("No query provided, skipping vector search")
                        vector_results = None
                except Exception as e:
                    logger.error(f"Error in vector search: {str(e)}")
                    vector_results = None

                # Process results
                bm25_docs = []
                vector_docs = []

                # Safely extract documents from BM25 results
                if bm25_results and isinstance(bm25_results, dict):
                    if "data" in bm25_results and isinstance(bm25_results["data"], dict):
                        if "Get" in bm25_results["data"] and isinstance(bm25_results["data"]["Get"], dict):
                            if "Document" in bm25_results["data"]["Get"] and isinstance(bm25_results["data"]["Get"]["Document"], list):
                                bm25_docs = bm25_results["data"]["Get"]["Document"]
                                if bm25_docs:  # Only extend if we have documents
                                    results.extend(bm25_docs)
                                    logger.info(f"Added {len(bm25_docs)} documents from BM25 search")

                # Safely extract documents from vector results
                if vector_results and isinstance(vector_results, dict):
                    if "data" in vector_results and isinstance(vector_results["data"], dict):
                        if "Get" in vector_results["data"] and isinstance(vector_results["data"]["Get"], dict):
                            if "Document" in vector_results["data"]["Get"] and isinstance(vector_results["data"]["Get"]["Document"], list):
                                vector_docs = vector_results["data"]["Get"]["Document"]
                                if vector_docs:  # Only extend if we have documents
                                    results.extend(vector_docs)
                                    logger.info(f"Added {len(vector_docs)} documents from vector search")

                # Break if no more results
                if (not bm25_docs and not vector_docs) or len(results) >= 10000:
                    logger.info("No more results found or reached maximum limit")
                    break

                offset += batch_size

            except Exception as e:
                logger.error(f"Error in search batch: {str(e)}")
                logger.error(f"Stack trace: {traceback.format_exc()}")
                break

        # Remove duplicates based on document_name and format results
        seen = set()
        unique_results = []
        for doc in results:
            if doc and isinstance(doc, dict) and "document_name" in doc and doc["document_name"] not in seen:
                seen.add(doc["document_name"])
                # Format the document metadata with default values
                formatted_doc = {
                    "document_name": doc.get("document_name", ""),
                    "content": doc.get("content", ""),
                    "case_title": doc.get("case_title", ""),
                    "citation": doc.get("citation", ""),
                    "court": doc.get("court", ""),
                    "jurisdiction": doc.get("jurisdiction", ""),
                    "decision_date": doc.get("decision_date", ""),
                    "calendar_date": doc.get("calendar_date", ""),
                    "docket_number": doc.get("docket_number", ""),
                    "parties": doc.get("parties", ""),
                    "judges": doc.get("judges", []),
                    "authoring_judge": doc.get("authoring_judge", ""),
                    "petitioner_attorney": doc.get("petitioner_attorney", ""),
                    "respondent_attorney": doc.get("respondent_attorney", ""),
                    "document_status": doc.get("document_status", ""),
                    "source": doc.get("source", ""),
                    "keywords": doc.get("keywords", []),
                    "statutes_cited": doc.get("statutes_cited", []),
                    "issue": doc.get("issue", ""),
                    "outcome": doc.get("outcome", ""),
                    "prior_history": doc.get("prior_history", "")
                }
                unique_results.append(formatted_doc)

        # Log the results for debugging
        logger.info(f"Found {len(unique_results)} unique results")
        for doc in unique_results:
            logger.info(f"Document: {doc['document_name']}")
            logger.info(f"Judges: {doc['judges']}")
            logger.info(f"Court: {doc['court']}")

        return {"results": unique_results}

    except Exception as e:
        logger.error(f"Error in search_court_cases: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return {"error": str(e)}, 500

@app.route('/court_cases')
@login_required
def court_cases():
    """Render the court cases search page"""
    return render_template('court_cases.html')

@app.route('/search_court_cases', methods=['POST'])
@login_required
def search_court_cases_route():
    """Handle court cases search requests"""
    try:
        data = request.get_json()
        query = data.get('query', '')
        filters = {
            'case_title': data.get('case_title', ''),
            'citation': data.get('citation', ''),
            'court': data.get('court', ''),
            'jurisdiction': data.get('jurisdiction', ''),
            'decision_date': data.get('decision_date', ''),
            'docket_number': data.get('docket_number', ''),
            'parties': data.get('parties', ''),
            'judges': data.get('judges', ''),
            'authoring_judge': data.get('authoring_judge', '')
        }
        
        if not query and not any(filters.values()):
            return jsonify({"error": "Please provide a search query or at least one filter"}), 400
            
        results = search_court_cases(query, filters)
        if isinstance(results, tuple) and len(results) == 2 and results[1] == 500:
            return jsonify({"error": results[0]}), 500
            
        # Ensure we're returning the results in the expected format
        if isinstance(results, dict) and "results" in results:
            return jsonify(results)
        else:
            return jsonify({"results": results})
            
    except Exception as e:
        logger.error(f"Error in search_court_cases_route: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500

def upload_court_case(title, date, state, verdict, description):
    """Upload a court case to Weaviate"""
    try:
        court_case = {
            "title": title,
            "date": date,
            "state": state,
            "verdict": verdict,
            "description": description
        }
        
        client.data_object.create(
            class_name="CourtCase",
            data_object=court_case
        )
        logger.info(f"Successfully uploaded court case: {title}")
        return True
    except Exception as e:
        logger.error(f"Error uploading court case: {str(e)}")
        return False

@app.route('/upload_court_case', methods=['POST'])
@login_required
def upload_court_case_route():
    """Handle court case upload requests"""
    try:
        data = request.get_json()
        title = data.get('title')
        date = data.get('date')
        state = data.get('state')
        verdict = data.get('verdict')
        description = data.get('description')
        
        if not all([title, date, state, verdict, description]):
            return jsonify({"error": "Missing required fields"}), 400
            
        if upload_court_case(title, date, state, verdict, description):
            return jsonify({"message": "Court case uploaded successfully"})
        else:
            return jsonify({"error": "Failed to upload court case"}), 500
    except Exception as e:
        logger.error(f"Error in court case upload: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/document_review')
@login_required
def document_review():
    """Render the document review page"""
    return render_template('document_review.html')

@app.route('/review_document', methods=['POST'])
@login_required
def review_document_route():
    """Handle document review requests"""
    try:
        if 'document' not in request.files:
            return jsonify({"error": "No document provided"}), 400
            
        file = request.files['document']
        doc_type = request.form.get('documentType')
        
        if not file or not doc_type:
            return jsonify({"error": "Missing required fields"}), 400
            
        # Create a temporary file with a unique name
        temp_dir = tempfile.gettempdir()
        temp_file_path = os.path.join(temp_dir, f"doc_review_{os.urandom(8).hex()}{os.path.splitext(file.filename)[1]}")
        
        try:
            # Save the uploaded file
            file.save(temp_file_path)
            
            # Review the document
            result = review_document(temp_file_path, doc_type)
            return jsonify(result)
            
        finally:
            # Ensure the file is closed and deleted
            try:
                if os.path.exists(temp_file_path):
                    # Try to close any open handles
                    import gc
                    gc.collect()
                    # Wait a short time to allow any processes to release the file
                    import time
                    time.sleep(0.1)
                    # Delete the file
                    os.unlink(temp_file_path)
            except Exception as e:
                logger.warning(f"Warning: Could not delete temporary file {temp_file_path}: {str(e)}")
                # Continue execution even if file deletion fails
                
    except Exception as e:
        logger.error(f"Error in document review: {str(e)}")
        return jsonify({"error": str(e)}), 500

def scan_document_for_fields(document_path, document_type):
    """
    Scan a document using Gemini AI to identify missing and filled fields.
    """
    try:
        logger.info(f"Starting document scan for: {document_path}")
        logger.info(f"Document type: {document_type}")
        
        # Extract text from the document
        text = extract_text_from_file(document_path)
        logger.info(f"Document text: {text}")
        
        if not text:
            logger.error("Could not extract text from the document")
            return {
                'missing_fields': [],
                'improper_fields': [],
                'error': 'Could not extract text from the document. The file might be encrypted or corrupted.'
            }
        
        # Prepare prompt for Gemini AI
        prompt = f"""
        Analyze this {document_type} document and identify:
        1. Which fields are missing (not present in the document)
        2. Which fields are present and filled by the user
        3. Which fields are present but empty
        
        Document Type: {document_type}
        Document Content:
        {text}
        
        Return the analysis in this exact JSON format:
        {{
            "missing_fields": ["list of missing field names"],
            "filled_fields": ["list of fields that are present and filled"],
            "empty_fields": ["list of fields that are present but empty"],
            "analysis": "brief explanation of the findings"
        }}
        """
        
        logger.info("Sending document to Gemini AI for analysis...")
        response = generate_gemini_response(prompt)
        
        if not response:
            logger.error("Failed to get response from Gemini AI")
            return {
                'missing_fields': [],
                'improper_fields': [],
                'error': 'Failed to analyze document with AI'
            }
        
        logger.info("Received response from Gemini AI")
        logger.debug(f"Raw AI response: {response}")
        
        try:
            # Extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if not json_match:
                logger.error("Could not find JSON in AI response")
                raise Exception("Could not find JSON in AI response")
            
            import json
            analysis = json.loads(json_match.group(0))
            
            # Log the analysis results
            logger.info("Analysis Results:")
            logger.info(f"Missing Fields: {analysis.get('missing_fields', [])}")
            logger.info(f"Filled Fields: {analysis.get('filled_fields', [])}")
            logger.info(f"Empty Fields: {analysis.get('empty_fields', [])}")
            logger.info(f"Analysis: {analysis.get('analysis', '')}")
            
            # Convert empty fields to improper fields format
            improper_fields = [{'field': field, 'reason': 'Field is present but empty'} 
                             for field in analysis.get('empty_fields', [])]
            
            return {
                'missing_fields': analysis.get('missing_fields', []),
                'improper_fields': improper_fields,
                'analysis': analysis.get('analysis', ''),
                'error': None
            }
            
        except Exception as e:
            logger.error(f"Error parsing AI response: {str(e)}")
            logger.error(f"Stack trace: {traceback.format_exc()}")
            return {
                'missing_fields': [],
                'improper_fields': [],
                'error': f'Error parsing AI response: {str(e)}'
            }
            
    except Exception as e:
        logger.error(f"Error scanning document: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return {
            'missing_fields': [],
            'improper_fields': [],
            'error': str(e)
        }

@app.route('/scan_document', methods=['POST'])
@login_required
def scan_document():
    try:
        if 'document' not in request.files:
            return jsonify({
                'status': 'error',
                'message': 'No document file provided'
            }), 400
            
        file = request.files['document']
        document_type = request.form.get('document_type')
        
        if not document_type:
            return jsonify({
                'status': 'error',
                'message': 'Document type not specified'
            }), 400
            
        if file.filename == '':
            return jsonify({
                'status': 'error',
                'message': 'No selected file'
            }), 400
            
        # Validate file extension
        allowed_extensions = {'pdf', 'doc', 'docx'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        if file_extension not in allowed_extensions:
            return jsonify({
                'status': 'error',
                'message': f'Invalid file type. Allowed types: {", ".join(allowed_extensions)}'
            }), 400
            
        # Save the uploaded file temporarily
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, secure_filename(file.filename))
        file.save(temp_path)
        
        try:
            # Scan the document
            scan_results = scan_document_for_fields(temp_path, document_type)
            document_text = extract_text_from_file(temp_path)
            if scan_results.get('error'):
                return jsonify({
                    'status': 'error',
                    'message': scan_results['error'],
                    'document_text': document_text or ''
                }), 400
            # Prepare response
            if not scan_results['missing_fields'] and not scan_results['improper_fields']:
                return jsonify({
                    'status': 'complete',
                    'message': 'Document scan completed successfully. All mandatory fields are properly filled.',
                    'missing_fields': [],
                    'improper_fields': [],
                    'analysis': scan_results.get('analysis', ''),
                    'document_text': document_text or ''
                })
            else:
                return jsonify({
                    'status': 'incomplete',
                    'message': 'Document scan completed. Some fields need attention.',
                    'missing_fields': scan_results['missing_fields'],
                    'improper_fields': scan_results['improper_fields'],
                    'analysis': scan_results.get('analysis', ''),
                    'document_text': document_text or ''
                })
                
        finally:
            # Clean up the temporary file
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except Exception as e:
                logger.warning(f"Warning: Could not delete temporary file {temp_path}: {str(e)}")
            
    except Exception as e:
        logger.error(f"Error processing document scan: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'An error occurred while scanning the document: {str(e)}'
        }), 500

@app.route('/document_scanner')
@login_required
def document_scanner():
    """Render the document scanner page"""
    return render_template('document_scanner.html')

@app.route('/test/documents')
@login_required
def test_documents():
    try:
        if client is None:
            return jsonify({"error": "Weaviate client is not initialized"}), 500
        
        # Build query to get all documents
        query = client.query.get("Document", ["document_name", "doc_type"]).with_additional("id")
        
        # Execute query
        result = query.do()
        
        if not result or "data" not in result or "Get" not in result["data"] or "Document" not in result["data"]["Get"]:
            logger.warning("No documents found in Weaviate")
            return jsonify({"error": "No documents found"}), 404
        
        # Log all document types found in the database
        logger.info("All document types found in database:")
        for obj in result["data"]["Get"]["Document"]:
            doc_name = obj["document_name"]
            doc_type = obj.get("doc_type", "Unknown")
            logger.info(f"Document: {doc_name}, Type: {doc_type}")
        
        # Process results
        documents = []
        logger.info("Processing documents from Weaviate:")
        for obj in result["data"]["Get"]["Document"]:
            doc_name = obj["document_name"]
            doc_type = obj.get("doc_type", "Unknown")
            logger.info(f"Found document: {doc_name} with type: {doc_type}")
            documents.append({
                "id": obj["_additional"]["id"],
                "name": doc_name,
                "type": doc_type
            })
        
        # Group documents by type
        documents_by_type = {}
        for doc in documents:
            doc_type = doc["type"]
            # Normalize document type
            doc_type = doc_type.strip()
            if doc_type.lower() == "legal templates":
                doc_type = "Legal templates"
            elif doc_type.lower() == "legal cases":
                doc_type = "Legal cases"
            elif doc_type.lower() == "texas constitution and statutes":
                doc_type = "Texas Constitution and Statutes"
            
            if doc_type not in documents_by_type:
                documents_by_type[doc_type] = []
            documents_by_type[doc_type].append(doc)
        
        # Sort documents by name within each type
        for doc_type in documents_by_type:
            documents_by_type[doc_type].sort(key=lambda x: x["name"])
        
        # Log the grouped documents
        logger.info("Documents grouped by type:")
        for doc_type, docs in documents_by_type.items():
            logger.info(f"Type: {doc_type} - Count: {len(docs)}")
            for doc in docs:
                logger.info(f"  - {doc['name']}")
        
        # Log all available document types
        logger.info(f"All available document types: {list(documents_by_type.keys())}")
        
        # Add empty lists for missing document types
        default_types = ["Legal templates", "Legal cases", "General", "Texas Constitution and Statutes", "FAQ"]
        for doc_type in default_types:
            if doc_type not in documents_by_type:
                documents_by_type[doc_type] = []
                logger.info(f"Added empty list for document type: {doc_type}")
        
        return render_template('test_documents.html', 
                             documents_by_type=documents_by_type,
                             total_documents=len(documents))
    except Exception as e:
        logger.error(f"Error in test_documents: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500

@app.route('/search')
@login_required
def search_documents():
    try:
        if client is None:
            logger.error("Weaviate client is not initialized")
            return render_template('search.html', 
                                documents=[], 
                                error="Weaviate connection is not available. Please check your configuration.",
                                current_page=1,
                                total_pages=1,
                                per_page=10,
                                query='',
                                doc_type='',
                                doc_types=[],
                                all_documents=[])

        # Get query parameters
        query = request.args.get('query', '')
        page = int(request.args.get('page', 1))
        per_page = int(request.args.get('per_page', 10))
        doc_type = request.args.get('doc_type', '')
        
        logger.info(f"Searching documents with query: {query}, type: {doc_type}")
        
        # Prepare type filter
        filter_clause = None
        if doc_type:
            doc_type = doc_type.replace('%20', ' ')
            logger.info(f"Searching for documents with type: {doc_type}")

            filter_clause = {
                        "path": ["doc_type"],
                        "operator": "Equal",
                "valueString": doc_type
            }

        # Fetch all documents using offset-based pagination
        offset = 0
        batch_size = 3000
        all_results = []

        while True:
            search_query = (
                client.query
                .get("Document", ["document_name", "doc_type", "content"])
                .with_additional("id")
                .with_limit(batch_size)
                .with_offset(offset)
            )

            if filter_clause:
                search_query = search_query.with_where(filter_clause)

            if query:
                search_query = search_query.with_near_text({
                    "concepts": [query],
                    "certainty": 0.7
                })

            result = search_query.do()
            documents = result.get("data", {}).get("Get", {}).get("Document", [])

            if not documents:
                break

            all_results.extend(documents)
            offset += batch_size

        # Process all documents to get unique document names and their counts
        doc_info = {}
        for obj in all_results:
            doc_name = obj.get("document_name", "Unknown")  # Handle None case
            doc_type_obj = obj.get("doc_type", "Unknown")
            content = obj.get("content", "")

            if doc_name not in doc_info:
                doc_info[doc_name] = {
                    "type": doc_type_obj,
                    "count": 0,
                    "content": content
                }
            doc_info[doc_name]["count"] += 1

        # Convert to list and sort
        all_documents = [{"name": name, "type": info["type"], "count": info["count"], "content": info["content"]} 
                         for name, info in doc_info.items()]
        all_documents.sort(key=lambda x: x["name"] or "")  # Handle None case in sorting

        # Format documents for dropdown - only include valid documents
        dropdown_documents = [
            {"name": doc["name"], "type": doc["type"]}
            for doc in all_documents
            if doc["name"] and doc["name"] != "Unknown"
        ]

        # Pagination logic
        total_count = len(all_documents)
        total_pages = (total_count + per_page - 1) // per_page
        start_idx = (page - 1) * per_page
        end_idx = min(start_idx + per_page, total_count)
        documents = all_documents[start_idx:end_idx]

        logger.info(f"Found {total_count} documents, showing {len(documents)} on page {page}")
        logger.info(f"Document types found: {set(doc['type'] for doc in all_documents)}")
        logger.info(f"Documents found: {[doc['name'] for doc in all_documents]}")

        # Populate dropdown
        doc_types = get_unique_document_types()
        logger.info(f"Available document types: {doc_types}")

        # Add max and min functions to template context
        template_context = {
            'documents': documents,
            'current_page': page,
            'total_pages': total_pages,
            'per_page': per_page,
            'query': query,
            'doc_type': doc_type,
            'doc_types': doc_types,
            'all_documents': dropdown_documents,  # Use formatted documents for dropdown
            'max': max,
            'min': min
        }

        return render_template('search.html', **template_context)
    except Exception as e:
        logger.error(f"Error searching documents: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return render_template(
            'search.html',
            documents=[],
            error=str(e),
            current_page=1,
            total_pages=1,
            per_page=10,
            query='',
            doc_type='',
            doc_types=[],
            all_documents=[],
            max=max,
            min=min
        )

@app.route('/get_documents')
@login_required
def get_documents():
    """Get documents filtered by document type"""
    try:
        if client is None:
            return jsonify({
                'error': 'Weaviate connection is not available. Please check your configuration.'
            }), 503

        doc_type = request.args.get('doc_type', '')
        logger.info(f"Fetching documents for type: {doc_type}")

        # Prepare type filter
        filter_clause = None
        if doc_type:
            doc_type = doc_type.replace('%20', ' ')
            logger.info(f"Searching for documents with type: {doc_type}")

            filter_clause = {
                        "path": ["doc_type"],
                        "operator": "Equal",
                "valueString": doc_type
            }

        # Fetch documents using offset-based pagination with increased batch size
        offset = 0
        batch_size = 10000  # Increased from 3000 to 10000
        all_results = []

        while True:
            query = (
                client.query
                .get("Document", ["document_name", "doc_type"])
                .with_additional("id")
                .with_limit(batch_size)
                .with_offset(offset)
            )

            if filter_clause:
                query = query.with_where(filter_clause)

            result = query.do()
            documents = result.get("data", {}).get("Get", {}).get("Document", [])

            if not documents:
                break

            all_results.extend(documents)
            offset += batch_size
            logger.info(f"Fetched {len(documents)} documents in current batch, total so far: {len(all_results)}")

        # Process results to get unique document names
        doc_info = {}
        for obj in all_results:
            doc_name = obj.get("document_name", "Unknown")
            doc_type_obj = obj.get("doc_type", "Unknown")

            if doc_name not in doc_info:
                doc_info[doc_name] = {
                    "type": doc_type_obj,
                    "count": 0
                }
            doc_info[doc_name]["count"] += 1

        # Convert to list and sort
        documents = [
            {"name": name, "type": info["type"], "count": info["count"]}
            for name, info in doc_info.items()
            if name and name != "Unknown"
        ]
        documents.sort(key=lambda x: x["name"] or "")

        logger.info(f"Found {len(documents)} documents for type {doc_type}")
        return jsonify({"documents": documents})

    except Exception as e:
        logger.error(f"Error fetching documents: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return jsonify({"error": str(e)}), 500

@app.route('/clients')
@login_required
def clients():
    search_query = request.args.get('search', '')
    if search_query:
        clients = Client.search(search_query)
    else:
        clients = Client.get_all()
    return render_template('clients.html', clients=clients)

@app.route('/clients/add')
@login_required
def add_client_page():
    return render_template('add_client.html')

@app.route('/client', methods=['POST'])
@login_required
def add_client():
    """Create a new client."""
    try:
        # Get form data
        first_name = request.form.get('first_name')
        last_name = request.form.get('last_name')
        email = request.form.get('email')
        phone_number = request.form.get('phone_number')
        alternate_number = request.form.get('alternate_number')
        address = request.form.get('address')
        city = request.form.get('city')
        state = request.form.get('state')
        zip_code = request.form.get('zip_code')
        country = request.form.get('country')
        client_type = request.form.get('client_type')
        preferred_contact = request.form.get('preferred_contact')

        # Log received data for debugging
        logger.info(f"Received client data: first_name={first_name}, last_name={last_name}, email={email}, phone={phone_number}, type={client_type}")

        # Validate required fields
        required_fields = {
            'first_name': first_name,
            'last_name': last_name,
            'email': email,
            'phone_number': phone_number,
            'client_type': client_type
        }
        
        missing_fields = [field for field, value in required_fields.items() if not value]
        if missing_fields:
            logger.warning(f"Missing required fields: {', '.join(missing_fields)}")
            return jsonify({'error': f'Missing required fields: {", ".join(missing_fields)}'}), 400

        # Check if email already exists
        existing_client = User.db.clients.find_one({'email': email})
        if existing_client:
            return jsonify({'error': 'A client with this email address already exists'}), 400

        # Create new client
        client = Client(
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone_number=phone_number,
            alternate_number=alternate_number,
            address=address,
            city=city,
            state=state,
            zip_code=zip_code,
            country=country,
            client_type=client_type,
            preferred_contact=preferred_contact
        )

        # Save to database
        client.save()

        return jsonify({
            'success': True,
            'message': 'Client created successfully',
            'client': client.to_dict()
        })

    except Exception as e:
        logger.error(f"Error creating client: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/client/<client_id>')
@login_required
def get_client(client_id):
    try:
        client = Client.get_by_id(client_id)
        if not client:
            return jsonify({'error': 'Client not found'}), 404
            
        # Convert client data to a dictionary format using to_dict()
        client_data = client.to_dict()
        
        return jsonify(client_data)
    except Exception as e:
        logger.error(f"Error fetching client: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 500

@app.route('/client/<client_id>', methods=['PUT'])
@login_required
def update_client(client_id):
    try:
        # Get the client first
        client = Client.get_by_id(client_id)
        if not client:
            return jsonify({'error': 'Client not found'}), 404

        # Get update data from request
        update_data = request.json

        # Update client attributes
        for key, value in update_data.items():
            if hasattr(client, key):
                setattr(client, key, value)

        # Update the client in the database
        client.update()

        return jsonify({'success': True, 'message': 'Client updated successfully'})
    except Exception as e:
        logger.error(f"Error updating client: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/client/<client_id>', methods=['DELETE'])
@login_required
def delete_client(client_id):
    try:
        result = Client.delete(client_id)
        if result.deleted_count > 0:
            return jsonify({'success': True})
        return jsonify({'error': 'Client not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/cases')
@login_required
def cases():
    """Display all cases."""
    try:
        # Get all cases using the Case model
        cases = Case.get_all()
    
        # Get all clients for the dropdown
        clients = Client.get_all()
        
        # Format the cases data
        formatted_cases = []
        for case in cases:
            try:
                # Get client information
                client = Client.get_by_id(case.get('client_id'))
                if client:
                    client_data = client.to_dict()
                    client_name = f"{client_data['first_name']} {client_data['last_name']}"
                else:
                    client_name = 'Unknown'
                
                formatted_case = {
                    '_id': str(case['_id']),
                    'case_number': case.get('case_number', 'N/A'),
                    'client_name': client_name,
                    'case_type': case.get('case_type', 'N/A'),
                    'status': case.get('status', 'N/A'),
                    'start_date': case.get('start_date', 'N/A'),
                    'priority': case.get('priority', 'N/A')
                }
                formatted_cases.append(formatted_case)
            except Exception as case_error:
                logger.error(f"Error formatting case {case.get('_id')}: {str(case_error)}")
                continue
        
        return render_template('cases.html', cases=formatted_cases, clients=clients)
    except Exception as e:
        logger.error(f"Error loading cases: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        flash('Error loading cases. Please try again.', 'error')
        return render_template('cases.html', cases=[], clients=[])

@app.route('/cases/add', methods=['GET', 'POST'])
@login_required
def add_case_page():
    if request.method == 'GET':
        clients = Client.get_all()
        selected_client_id = request.args.get('client_id')
        return render_template('add_case.html', clients=clients, selected_client_id=selected_client_id)
    
    try:
        # Log the incoming request data
        logger.debug("Request form data: %s", request.form)
        logger.debug("Request files: %s", request.files)
        
        # Get form data
        case_data = {
            'client_id': request.form.get('client_id'),
            'title': request.form.get('title', ''),
            'description': request.form.get('description', ''),
            'case_type': request.form.get('case_type'),
            'court_name': request.form.get('court_name', ''),
            'case_number': request.form.get('case_number'),
            'status': request.form.get('status'),
            'start_date': request.form.get('start_date'),
            'end_date': request.form.get('end_date') if request.form.get('end_date') else None,
            'priority': request.form.get('priority'),
            'created_at': datetime.utcnow(),
            'updated_at': datetime.utcnow()
        }
        
        # Validate required fields
        required_fields = ['client_id', 'title', 'case_type', 'case_number', 'status', 'start_date', 'priority']
        missing_fields = [field for field in required_fields if not case_data.get(field)]
        if missing_fields:
            flash(f'Missing required fields: {", ".join(missing_fields)}', 'error')
            return redirect(url_for('add_case_page'))
        
        # Get client data using client_id (UUID)
        client = Client.get_by_id(case_data['client_id'])
        if not client:
            flash('Client not found', 'error')
            return redirect(url_for('add_case_page'))
            
        # Add client name to case data
        client_data = client.to_dict()
        case_data['client_name'] = f"{client_data['first_name']} {client_data['last_name']}"
        
        # Create case using the Case model
        case = Case(**case_data)
        case.save()
        
        flash('Case added successfully!', 'success')
        return redirect(url_for('cases'))
        
    except Exception as e:
        logger.error(f"Error adding case: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        flash(f'Error adding case: {str(e)}', 'error')
        return redirect(url_for('add_case_page'))

@app.route('/case/<case_id>')
@login_required
def get_case(case_id):
    try:
        case = Case.get_by_id(case_id)
        if not case:
            return jsonify({'error': 'Case not found'}), 404
            
        # Get client information using client_id
        client = Client.get_by_id(case.client_id)
        if client:
            client_data = client.to_dict()
            case.client_name = f"{client_data['first_name']} {client_data['last_name']}"
        else:
            case.client_name = 'Unknown Client'
        
        # Convert case object to dictionary
        case_dict = {
            '_id': case._id,
            'client_id': case.client_id,
            'client_name': case.client_name,
            'title': case.title,
            'description': case.description,
            'case_type': case.case_type,
            'court_name': case.court_name,
            'case_number': case.case_number,
            'status': case.status,
            'priority': case.priority,
            'documents': case.documents,
            'court_visits': case.court_visits,
            'created_at': case.created_at,
            'updated_at': case.updated_at
        }
        
        # Format dates
        if case.start_date:
            case_dict['start_date'] = case.start_date.strftime('%Y-%m-%d')
        if case.end_date:
            case_dict['end_date'] = case.end_date.strftime('%Y-%m-%d')
            
        return jsonify(case_dict)
    except Exception as e:
        logger.error(f"Error fetching case: {str(e)}")
        logger.error(f"Stack trace: {traceback.format_exc()}")
        return jsonify({'error': str(e)}), 400

@app.route('/case/<case_id>', methods=['PUT'])
@login_required
def update_case(case_id):
    logger.debug(f"Accessing edit case page for case ID: {case_id}")
    try:
        case = Case.get_by_id(case_id)
        if not case:
            logger.error(f"Case not found with ID: {case_id}")
            return jsonify({'error': 'Case not found'}), 404

        if request.method == 'POST':
            logger.debug(f"Processing edit case POST request for case ID: {case_id}")
            try:
                # Update case data
                case.title = request.form.get('title')
                case.description = request.form.get('description')
                case.case_type = request.form.get('case_type')
                case.court_name = request.form.get('court_name')
                case.case_number = request.form.get('case_number')
                case.status = request.form.get('status')
                case.priority = request.form.get('priority')
                case.client_id = request.form.get('client_id')
                
                # Handle dates
                start_date = request.form.get('start_date')
                if start_date:
                    case.start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = request.form.get('end_date')
                if end_date:
                    case.end_date = datetime.strptime(end_date, '%Y-%m-%d')
                
                logger.debug(f"Updating case with new data - Title: {case.title}, Status: {case.status}")
                case.save()
                logger.debug("Case updated successfully")
                
                flash('Case updated successfully!', 'success')
                return redirect(url_for('cases'))
            except Exception as e:
                logger.error(f"Error updating case: {str(e)}")
                flash(f'Error updating case: {str(e)}', 'error')
                return redirect(url_for('edit_case', case_id=case_id))

        # Get all clients for the dropdown
        clients = Client.get_all()
        logger.debug("Rendering edit case form")
        return render_template('edit_case.html', case=case, clients=clients)
    except Exception as e:
        logger.error(f"Error in edit case route: {str(e)}")
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('cases'))

@app.route('/case/<case_id>', methods=['DELETE'])
@login_required
def delete_case(case_id):
    case = Case.get_by_id(case_id)
    if not case:
        return jsonify({'error': 'Case not found'}), 404
    
    try:
        Case.delete(case_id)
        return jsonify({'message': 'Case deleted successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/client/<client_id>/cases')
@login_required
def get_client_cases(client_id):
    try:
        cases = Case.get_by_client_id(client_id)
        return jsonify(cases)
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/appointments')
@login_required
def appointments():
    try:
        logger.debug("Starting appointments route")
        # Get all appointments and sort by date_time in descending order
        logger.debug("Fetching all appointments")
        appointments = Appointment.get_all()
        appointments.sort(key=lambda x: x.get('date_time', datetime.min), reverse=True)
        logger.debug(f"Retrieved and sorted {len(appointments)} appointments")
        # Get all cases for the dropdown
        logger.debug("Fetching all cases")
        cases = Case.get_all()
        logger.debug(f"Retrieved {len(cases)} cases")
        formatted_cases = []
        for case in cases:
            try:
                case_id = case['_id']
                client_id = case.get('client_id')
                # Get client information
                client = Client.get_by_id(client_id)
                if client:
                    client_data = client.to_dict()
                    case_title = f"{case.get('case_number', '')} - {client_data.get('first_name', '')} {client_data.get('last_name', '')}"
                else:
                    case_title = f"{case.get('case_number', '')} - Unknown Client"
                formatted_case = {
                    '_id': str(case_id),
                    'title': case_title
                }
                formatted_cases.append(formatted_case)
            except Exception as case_error:
                logger.error(f"Error formatting case {case_id if 'case_id' in locals() else 'unknown'}: {str(case_error)}")
                logger.error(f"Case data: {case}")
                continue
        logger.debug(f"Formatted cases for dropdown: {formatted_cases}")
        logger.debug(f"Total formatted cases: {len(formatted_cases)}")
        # Format appointments with case information
        formatted_appointments = []
        for appointment in appointments:
            try:
                appointment_id = appointment.get('appointment_id')
                logger.debug(f"Processing appointment: {appointment_id}")
        # Get case information
                case = Case.get_by_id(appointment.get('case_id'))
                if case:
                    case_id = case._id
                    logger.debug(f"Found case for appointment {appointment_id}: {case_id}")
                    # Get client information
                    client = Client.get_by_id(case.client_id)
                    if client:
                        client_data = client.to_dict()
                        logger.debug(f"Found client for case {case_id}: {client_data.get('_id')}")
                        case_title = f"{case.case_number} - {client_data['first_name']} {client_data['last_name']}"
                    else:
                        logger.debug(f"No client found for case {case_id}")
                        case_title = f"{case.case_number} - Unknown Client"
                else:
                    logger.debug(f"No case found for appointment {appointment_id}")
                    case_title = "Unknown Case"

                formatted_appointment = {
                    'appointment_id': appointment_id,
                    'case_id': appointment.get('case_id'),
                    'case_title': case_title,
                    'date_time': appointment.get('date_time'),
                    'location': appointment.get('location'),
                    'purpose': appointment.get('purpose'),
                    'status': appointment.get('status')
                }
                formatted_appointments.append(formatted_appointment)
            except Exception as appt_error:
                logger.error(f"Error formatting appointment {appointment.get('appointment_id')}: {str(appt_error)}")
                logger.error(f"Appointment data: {appointment}")
                continue

        logger.debug(f"Successfully formatted {len(formatted_appointments)} appointments and {len(formatted_cases)} cases")
        return render_template('appointments.html', 
                             appointments=formatted_appointments,
                             cases=formatted_cases)
    except Exception as e:
        logger.error(f"Error loading appointments: {str(e)}")
        logger.error(f"Full error details:", exc_info=True)
        flash('Error loading appointments. Please try again.', 'error')
        return render_template('appointments.html', appointments=[], cases=[])

@app.route('/appointment/<appointment_id>', methods=['PUT'])
@login_required
def update_appointment(appointment_id):
    try:
        update_data = request.json
        if 'date_time' in update_data:
            update_data['date_time'] = datetime.strptime(update_data['date_time'], '%Y-%m-%dT%H:%M')
        
        appointment = Appointment.update(appointment_id, update_data)
        return jsonify({'success': True, 'appointment': appointment})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/appointment/<appointment_id>', methods=['DELETE'])
@login_required
def delete_appointment(appointment_id):
    try:
        Appointment.delete(appointment_id)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/initial_inquiry')
@login_required
def initial_inquiry_page():
    """Display the initial inquiry page with recent inquiries."""
    # Create the collection if it doesn't exist
    InitialInquiry.create_collection()
    
    # Get recent inquiries
    inquiries = InitialInquiry.get_recent(5)  # Get 5 most recent inquiries
    
    return render_template('initial_inquiry.html', inquiries=inquiries)

@app.route('/add_initial_inquiry')
@login_required
def add_initial_inquiry_page():
    """Display the form for adding a new initial inquiry."""
    # Get list of attorneys for the dropdown
    attorneys = Attorney.get_all()
    return render_template('add_initial_inquiry.html', attorneys=attorneys)

@app.route('/create_initial_inquiry', methods=['POST'])
@login_required
def create_initial_inquiry():
    try:
        # Create new inquiry
        inquiry = InitialInquiry(
            fullName=request.form.get('fullName'),
            phoneNumber=request.form.get('phoneNumber'),
            email=request.form.get('email'),
            appointmentType=request.form.get('appointmentType'),
            preferredDate=request.form.get('preferredDate'),
            preferredTime=request.form.get('preferredTime'),
            referralSource=request.form.get('referralSource'),
            assignedAttorney=request.form.get('assignedAttorney'),
            caseDescription=request.form.get('caseDescription'),
            notes=request.form.get('notes')
        )
        inquiry.save()
        
        # Create appointment if preferred date and time are provided
        if inquiry.preferredDate and inquiry.preferredTime:
            appointment = Appointment(
                case_id=None,  # No case yet
                date_time=datetime.strptime(f"{inquiry.preferredDate} {inquiry.preferredTime}", "%Y-%m-%d %H:%M"),
                location="Office",
                purpose=f"Initial {inquiry.appointmentType}: {inquiry.caseDescription[:100]}",
                status="Scheduled"
            )
            appointment.save()
        
        flash('Initial inquiry created successfully!', 'success')
        return redirect(url_for('initial_inquiry_page'))
    except Exception as e:
        flash(f'Error creating initial inquiry: {str(e)}', 'error')
        return redirect(url_for('initial_inquiry_page'))

@app.route('/initial_inquiry/<inquiry_id>', methods=['GET'])
@login_required
def get_inquiry(inquiry_id):
    try:
        inquiry = InitialInquiry.get_by_id(inquiry_id)
        if inquiry:
            return jsonify(inquiry)
        return jsonify({'error': 'Inquiry not found'}), 404
    except Exception as e:
        logger.error(f"Error getting inquiry {inquiry_id}: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/initial_inquiry/<inquiry_id>', methods=['PUT'])
@login_required
def update_inquiry(inquiry_id):
    try:
        data = request.get_json()
        if InitialInquiry.update(inquiry_id, data):
            return jsonify({'message': 'Inquiry updated successfully'})
        return jsonify({'error': 'Inquiry not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/initial_inquiry/<inquiry_id>', methods=['DELETE'])
@login_required
def delete_inquiry(inquiry_id):
    try:
        if InitialInquiry.delete(inquiry_id):
            return jsonify({'message': 'Inquiry deleted successfully'})
        return jsonify({'error': 'Inquiry not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 400

@app.route('/client/<client_id>/edit')
@login_required
def edit_client_page(client_id):
    client = Client.get_by_id(client_id)
    if client:
        return render_template('edit_client.html', client=client)
    flash('Client not found', 'error')
    return redirect(url_for('clients'))

@app.route('/case/<case_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_case(case_id):
    logger.debug(f"Accessing edit case page for case ID: {case_id}")
    try:
        case = Case.get_by_id(case_id)
        if not case:
            logger.error(f"Case not found with ID: {case_id}")
            flash('Case not found', 'error')
            return redirect(url_for('cases'))

        if request.method == 'POST':
            logger.debug(f"Processing edit case POST request for case ID: {case_id}")
            try:
                # Update case data
                case.title = request.form.get('title')
                case.description = request.form.get('description')
                case.case_type = request.form.get('case_type')
                case.court_name = request.form.get('court_name')
                case.case_number = request.form.get('case_number')
                case.status = request.form.get('status')
                case.priority = request.form.get('priority')
                case.client_id = request.form.get('client_id')
                
                # Handle dates
                start_date = request.form.get('start_date')
                if start_date:
                    case.start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = request.form.get('end_date')
                if end_date:
                    case.end_date = datetime.strptime(end_date, '%Y-%m-%d')
                
                logger.debug(f"Updating case with new data - Title: {case.title}, Status: {case.status}")
                case.save()
                logger.debug("Case updated successfully")
                
                flash('Case updated successfully!', 'success')
                return redirect(url_for('cases'))
            except Exception as e:
                logger.error(f"Error updating case: {str(e)}")
                flash(f'Error updating case: {str(e)}', 'error')
                return redirect(url_for('edit_case', case_id=case_id))

        # Get all clients for the dropdown
        clients = Client.get_all()
        logger.debug("Rendering edit case form")
        return render_template('edit_case.html', case=case, clients=clients)
    except Exception as e:
        logger.error(f"Error in edit case route: {str(e)}")
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('cases'))

@app.route('/appointment/<appointment_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_appointment(appointment_id):
    try:
        logger.debug(f"Accessing edit appointment page for appointment ID: {appointment_id}")
        appointment = Appointment.get_by_id(appointment_id)
        
        if not appointment:
            logger.error(f"Appointment not found with ID: {appointment_id}")
            flash('Appointment not found', 'error')
            return redirect(url_for('appointments'))
        
        if request.method == 'POST':
            try:
                # Get form data
                case_id = request.form.get('case_id')  # This can be empty
                date_time = datetime.strptime(request.form['date_time'], '%Y-%m-%dT%H:%M')
                location = request.form['location']
                purpose = request.form['purpose']
                status = request.form['status']
                
                # Update appointment data
                update_data = {
                    'case_id': case_id,  # Can be None
                    'date_time': date_time,
                    'location': location,
                    'purpose': purpose,
                    'status': status
                }
                
                logger.debug(f"Updating appointment with data: {update_data}")
                Appointment.update(appointment_id, update_data)
                flash('Appointment updated successfully!', 'success')
                return redirect(url_for('appointments'))
                
            except ValueError as ve:
                logger.error(f"Invalid date format: {str(ve)}")
                flash('Invalid date format. Please use the correct format.', 'error')
                return redirect(url_for('edit_appointment', appointment_id=appointment_id))
            except Exception as e:
                logger.error(f"Error updating appointment: {str(e)}")
                flash(f'Error updating appointment: {str(e)}', 'error')
                return redirect(url_for('edit_appointment', appointment_id=appointment_id))
        
        # Get all cases for the dropdown
        cases = Case.get_all()
        formatted_cases = []
        for case in cases:
            try:
                client = Client.get_by_id(case.get('client_id'))
                if client:
                    client_data = client.to_dict()
                    case_title = f"{case.get('case_number')} - {client_data['first_name']} {client_data['last_name']}"
                else:
                    case_title = f"{case.get('case_number')} - Unknown Client"
                
                formatted_case = {
                    'case_id': str(case.get('_id')),
                    'title': case_title
                }
                formatted_cases.append(formatted_case)
            except Exception as case_error:
                logger.error(f"Error formatting case: {str(case_error)}")
                logger.error(f"Case data: {case}")
                continue
        
        logger.debug("Rendering edit appointment form")
        return render_template('edit_appointment.html', 
                             appointment=appointment, 
                             cases=formatted_cases)
                             
    except Exception as e:
        logger.error(f"Error in edit appointment route: {str(e)}")
        logger.error("Full error details:", exc_info=True)
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('appointments'))

@app.route('/case/<case_id>/documents', methods=['POST'])
@login_required
def upload_case_documents(case_id):
    logger.debug(f"Processing document upload for case ID: {case_id}")
    try:
        # Get case data from database
        case = Case.get_by_id(case_id)
        if not case:
            logger.error(f"Case not found with ID: {case_id}")
            return jsonify({'error': 'Case not found'}), 404

        logger.debug("Case object retrieved successfully")

        if 'files' not in request.files:
            logger.warning("No files provided in request")
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        description = request.form.get('description', '')
        logger.debug(f"Received {len(files)} files for upload")

        uploaded_documents = []
        for file in files:
            if file.filename:
                logger.debug(f"Processing file: {file.filename}")
                document = case.add_document(file, description)
                uploaded_documents.append(document)
                logger.debug(f"File {file.filename} uploaded successfully")

        logger.debug(f"Successfully uploaded {len(uploaded_documents)} documents")
        return jsonify({
            'message': 'Documents uploaded successfully',
            'documents': uploaded_documents
        })

    except Exception as e:
        logger.error(f"Error uploading documents: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/case/<case_id>/documents/<document_id>', methods=['DELETE'])
@login_required
def delete_case_document(case_id, document_id):
    logger.debug(f"Processing document deletion - Case ID: {case_id}, Document ID: {document_id}")
    try:
        case = Case.get_by_id(case_id)
        if not case:
            logger.error(f"Case not found with ID: {case_id}")
            return jsonify({'error': 'Case not found'}), 404

        logger.debug("Deleting document from case")
        case.delete_document(document_id)
        logger.debug("Document deleted successfully")
        return jsonify({'message': 'Document deleted successfully'})

    except Exception as e:
        logger.error(f"Error deleting document: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/document/<document_id>/download')
@login_required
def download_document(document_id):
    try:
        # Find the case containing this document
        db = get_db()
        case = db.cases.find_one({'documents._id': document_id})
        if not case:
            return jsonify({'error': 'Document not found'}), 404

        # Find the document
        document = next((doc for doc in case['documents'] if doc['_id'] == document_id), None)
        if not document:
            return jsonify({'error': 'Document not found'}), 404

        # Check if file exists
        if not os.path.exists(document['file_path']):
            return jsonify({'error': 'File not found'}), 404

        return send_file(
            document['file_path'],
            as_attachment=True,
            download_name=document['original_filename'],
            mimetype=document['mime_type']
        )

    except Exception as e:
        logger.error(f"Error downloading document: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/documents/upload', methods=['POST'])
@login_required
def upload_documents():
    logger.debug("Processing temporary document upload")
    try:
        if 'files' not in request.files:
            logger.warning("No files provided in request")
            return jsonify({'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        description = request.form.get('description', '')
        
        if not files:
            logger.warning("No files selected")
            return jsonify({'error': 'No files selected'}), 400
            
        # Get document storage path
        base_path = os.getenv('DOCUMENT_STORAGE_PATH', 'documents')
        if not os.path.exists(base_path):
            logger.debug(f"Creating document storage directory: {base_path}")
            os.makedirs(base_path)
            
        uploaded_documents = []
        for file in files:
            if file and file.filename:
                logger.debug(f"Processing file: {file.filename}")
                # Create a temporary case ID for document storage
                temp_case_id = str(ObjectId())
                case_path = os.path.join(base_path, temp_case_id)
                if not os.path.exists(case_path):
                    logger.debug(f"Creating case directory: {case_path}")
                    os.makedirs(case_path)
                
                # Save the file in version 1 directory
                version_dir = os.path.join(case_path, "v1")
                if not os.path.exists(version_dir):
                    logger.debug(f"Creating version directory: {version_dir}")
                    os.makedirs(version_dir)
                
                filename = secure_filename(file.filename)
                file_path = os.path.join(version_dir, filename)
                logger.debug(f"Saving file to: {file_path}")
                file.save(file_path)
                
                # Create document record
                document = {
                    '_id': str(ObjectId()),
                    'original_filename': file.filename,
                    'filename': filename,
                    'file_path': file_path,
                    'description': description,
                    'version': 1,
                    'upload_date': datetime.utcnow(),
                    'case_id': None,  # Will be set when case is created
                    'temp_case_id': temp_case_id  # Store temp case ID for cleanup
                }
                
                # Insert into MongoDB
                logger.debug("Inserting document record into database")
                result = documents_collection.insert_one(document)
                document['_id'] = str(result.inserted_id)
                uploaded_documents.append(document)
                logger.debug(f"File {file.filename} uploaded successfully")
        
        logger.debug(f"Successfully uploaded {len(uploaded_documents)} documents")
        return jsonify({'documents': uploaded_documents})
    except Exception as e:
        logger.error(f"Error uploading documents: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/case/documents/<document_id>', methods=['DELETE'])
@login_required
def delete_temp_document(document_id):
    try:
        # Find the document
        document = documents_collection.find_one({'_id': ObjectId(document_id)})
        if not document:
            return jsonify({'error': 'Document not found'}), 404

        # Delete the file
        if os.path.exists(document['file_path']):
            os.remove(document['file_path'])
            
            # Try to remove the version directory if it's empty
            version_dir = os.path.dirname(document['file_path'])
            if os.path.exists(version_dir) and not os.listdir(version_dir):
                os.rmdir(version_dir)
                
            # Try to remove the case directory if it's empty
            case_dir = os.path.dirname(version_dir)
            if os.path.exists(case_dir) and not os.listdir(case_dir):
                os.rmdir(case_dir)

        # Delete from MongoDB
        documents_collection.delete_one({'_id': ObjectId(document_id)})
        
        return jsonify({'message': 'Document deleted successfully'})
    except Exception as e:
        app.logger.error(f"Error deleting document: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/case/documents/temp_upload', methods=['POST'])
@login_required
def upload_temp_documents():
    """Handle temporary document uploads for new cases before they are created."""
    logger.debug("Processing temporary document upload for new case")
    try:
        if 'files' not in request.files:
            logger.warning("No files provided in request")
            return jsonify({'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        description = request.form.get('description', '')
        logger.debug(f"Received {len(files)} files for upload")
        
        # Get document storage path
        base_path = os.getenv('DOCUMENT_STORAGE_PATH', 'documents')
        if not os.path.exists(base_path):
            logger.debug(f"Creating document storage directory: {base_path}")
            os.makedirs(base_path)
            
        uploaded_documents = []
        for file in files:
            if file and file.filename:
                logger.debug(f"Processing file: {file.filename}")
                # Create a temporary case ID for document storage
                temp_case_id = str(ObjectId())
                case_path = os.path.join(base_path, temp_case_id)
                if not os.path.exists(case_path):
                    logger.debug(f"Creating case directory: {case_path}")
                    os.makedirs(case_path)
                
                # Save the file in version 1 directory
                version_dir = os.path.join(case_path, "v1")
                if not os.path.exists(version_dir):
                    logger.debug(f"Creating version directory: {version_dir}")
                    os.makedirs(version_dir)
                
                filename = secure_filename(file.filename)
                file_path = os.path.join(version_dir, filename)
                logger.debug(f"Saving file to: {file_path}")
                file.save(file_path)
                
                # Create document info without database insertion
                document = {
                    '_id': str(ObjectId()),  # Generate ID for frontend reference
                    'original_filename': file.filename,
                    'filename': filename,
                    'file_path': file_path,
                    'description': description,
                    'version': 1,
                    'temp_case_id': temp_case_id
                }
                
                uploaded_documents.append(document)
                logger.debug(f"File {file.filename} uploaded successfully")
        
        logger.debug(f"Successfully uploaded {len(uploaded_documents)} documents")
        return jsonify({
            'message': 'Documents uploaded successfully',
            'documents': uploaded_documents
        })
    except Exception as e:
        logger.error(f"Error uploading documents: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/case/<case_id>/court_visit', methods=['POST'])
@login_required
def add_court_visit(case_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        date = data.get('date')
        notes = data.get('notes')

        if not date or not notes:
            return jsonify({'error': 'Date and notes are required'}), 400

        # Get case
        case = Case.get_by_id(case_id)
        if not case:
            return jsonify({'error': 'Case not found'}), 404

        # Add court visit
        case.add_court_visit(date, notes)
        
        # Save the updated case
        case.save()

        return jsonify({
            'success': True,
            'message': 'Court visit added successfully',
            'visit': case.court_visits[-1]  # Return the newly added visit
        })

    except Exception as e:
        logger.error(f"Error adding court visit: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/case/<case_id>/court_visit/<visit_id>', methods=['GET'])
@login_required
def get_court_visit(case_id, visit_id):
    try:
        case = Case.get_by_id(case_id)
        if not case:
            return jsonify({'error': 'Case not found'}), 404

        visit = next((v for v in case.court_visits if v['_id'] == visit_id), None)
        if not visit:
            return jsonify({'error': 'Court visit not found'}), 404

        return jsonify({
            'date': visit['date'].strftime('%Y-%m-%d'),
            'notes': visit['notes']
        })
    except Exception as e:
        logger.error(f"Error getting court visit: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/case/<case_id>/court_visit/<visit_id>', methods=['PUT'])
@login_required
def update_court_visit(case_id, visit_id):
    try:
        data = request.get_json()
        if not data or 'date' not in data or 'notes' not in data:
            return jsonify({'error': 'Missing required fields'}), 400

        case = Case.get_by_id(case_id)
        if not case:
            return jsonify({'error': 'Case not found'}), 404

        # Convert date string to datetime
        visit_date = datetime.strptime(data['date'], '%Y-%m-%d')
        
        # Update court visit
        visit = case.update_court_visit(visit_id, data['notes'])
        if not visit:
            return jsonify({'error': 'Court visit not found'}), 404

        case.save()
        return jsonify({
            'success': True,
            'message': 'Court visit updated successfully',
            'visit': visit
        })
    except Exception as e:
        logger.error(f"Error updating court visit: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/case/<case_id>/court_visit/<visit_id>', methods=['DELETE'])
@login_required
def delete_court_visit(case_id, visit_id):
    try:
        case = Case.get_by_id(case_id)
        if not case:
            return jsonify({'error': 'Case not found'}), 404

        case.delete_court_visit(visit_id)
        case.save()

        return jsonify({
            'success': True,
            'message': 'Court visit deleted successfully'
        })
    except Exception as e:
        logger.error(f"Error deleting court visit: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/appointment/<appointment_id>', methods=['GET'])
@login_required
def get_appointment(appointment_id):
    try:
        appointment = Appointment.get_by_id(appointment_id)
        if not appointment:
            return jsonify({'error': 'Appointment not found'}), 404
            
        # Get case information if appointment has a case
        case_title = "No Case Assigned"
        if appointment.get('case_id'):
            case = Case.get_by_id(appointment.get('case_id'))
            if case:
                client = Client.get_by_id(case.client_id)
                if client:
                    client_data = client.to_dict()
                    case_title = f"{case.case_number} - {client_data['first_name']} {client_data['last_name']}"
                else:
                    case_title = f"{case.case_number} - Unknown Client"
            else:
                case_title = "Unknown Case"
        
        # Format appointment data
        appointment_data = {
            'appointment_id': appointment.get('appointment_id'),
            'case_id': appointment.get('case_id'),
            'case_title': case_title,
            'date_time': appointment.get('date_time').strftime('%Y-%m-%dT%H:%M') if appointment.get('date_time') else None,
            'location': appointment.get('location'),
            'purpose': appointment.get('purpose'),
            'status': appointment.get('status')
        }
        
        return jsonify(appointment_data)
    except Exception as e:
        logger.error(f"Error getting appointment: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/appointment', methods=['POST'])
@login_required
def add_appointment():
    try:
        case_id = request.form.get('case_id')
        date_time_str = request.form.get('date_time')
        location = request.form.get('location')
        purpose = request.form.get('purpose')
        status = request.form.get('status', 'Scheduled')

        # Validate required fields
        if not date_time_str or not location or not purpose or not status:
            return jsonify({'success': False, 'error': 'Missing required fields'}), 400

        # Parse date_time
        try:
            date_time = datetime.strptime(date_time_str, '%Y-%m-%dT%H:%M')
        except Exception:
            return jsonify({'success': False, 'error': 'Invalid date/time format'}), 400

        # Create and save the appointment
        appointment = Appointment(
            case_id=case_id,
            date_time=date_time,
            location=location,
            purpose=purpose,
            status=status
        )
        appointment.save()
        return jsonify({'success': True, 'appointment_id': appointment.appointment_id})
    except Exception as e:
        logger.error(f'Error adding appointment: {str(e)}')
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/upcoming-appointments')
@login_required
def api_upcoming_appointments():
    try:
        now = datetime.utcnow()
        two_weeks_later = now + timedelta(days=14)
        appointments = Appointment.get_all()
        upcoming = []
        for appt in appointments:
            dt = appt.get('date_time')
            if isinstance(dt, str):
                try:
                    dt = datetime.fromisoformat(dt)
                except Exception:
                    continue
            if not dt or not (now <= dt <= two_weeks_later):
                continue
            # Get case info
            case_title = None
            client_name = None
            if appt.get('case_id'):
                case = Case.get_by_id(appt.get('case_id'))
                if case:
                    case_title = getattr(case, 'case_number', None) or str(getattr(case, '_id', ''))
                    client_id = getattr(case, 'client_id', None)
                    client = Client.get_by_id(client_id) if client_id else None
                    if client:
                        client_name = f"{getattr(client, 'first_name', '')} {getattr(client, 'last_name', '')}".strip()
            upcoming.append({
                'appointment_id': appt.get('appointment_id'),
                'date_time': dt.strftime('%Y-%m-%d %H:%M'),
                'case_title': case_title,
                'client_name': client_name,
            })
        return jsonify({'appointments': upcoming, 'count': len(upcoming)})
    except Exception as e:
        logger.error(f"Error fetching upcoming appointments: {str(e)}")
        return jsonify({'appointments': [], 'count': 0, 'error': str(e)}), 500

@app.route('/case_analyzer')
@login_required
def case_analyzer():
    return render_template('case_analyzer.html')

@app.route('/api/case_analyzer', methods=['POST'])
@login_required
def api_case_analyzer():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        ext = file.filename.rsplit('.', 1)[-1].lower()
        temp_path = os.path.join(tempfile.gettempdir(), secure_filename(file.filename))
        file.save(temp_path)
        # Extract text
        if ext == 'pdf':
            text = extract_text_from_pdf(temp_path)
        elif ext in ('docx', 'doc'):
            text = extract_text_from_file(temp_path)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400
        if not text:
            return jsonify({'error': 'Could not extract text from file'}), 400
        # Local vector DB search (reuse perform_search or similar)
        local_results = perform_search('analyze this case', None)
        # Web search supplement (simulate with web_search or placeholder)
        web_supplement = 'Web search results: [Simulated trusted legal info]'
        # AI analysis prompt
        prompt = f"""
        You are a legal analysis assistant.

        Analyze the following court case document and return a structured JSON containing:

        - key_issues: List the central legal issues in the case.
        - weaknesses: Highlight any weaknesses in the arguments or evidence presented.
        - missing_documents: List any documentation that appears to be missing or insufficient.
        - suggested_strategies: Provide legal strategies to strengthen the case and prepare for court.
        - likely_hearing_questions: Suggest common or challenging questions that might arise in the next hearing.
        - winning_tips: Offer general AI-generated recommendations to improve the chances of winning.
        - case_facts: Summarize the critical facts relevant to the court.
        - legal_risks: Identify any potential legal vulnerabilities or oversights.
        - recommended_next_steps: Concrete actions the legal team or plaintiff should take.
        - jurisdiction_considerations: Mention any legal specifics based on the applicable jurisdiction.

        Use the following as context for your legal reasoning:
        {local_results}

        Additional knowledge from trusted sources:
        {web_supplement}

        Document Content:
        {text}

        Return ONLY valid, parseable JSON under a top-level key: "case_analysis_results"
        """
        ai_response = generate_gemini_response(prompt)
        import re, json
        json_match = re.search(r'\{.*\}', ai_response, re.DOTALL)
        if json_match:
            analysis = json.loads(json_match.group(0))
            if 'case_analysis_results' in analysis:
                analysis = analysis['case_analysis_results']
        else:
            analysis = {'raw': ai_response}
        return jsonify({'success': True, 'analysis': analysis})
    except Exception as e:
        logger.error(f'Error in case analyzer: {str(e)}')
        return jsonify({'success': False, 'error': str(e)})

if __name__ == '__main__':
    try:
        create_schema()
        logger.info("Starting application...")
        # Create temp directory if it doesn't exist
        os.makedirs('temp', exist_ok=True)
        logger.info("Starting Flask server")
        app.run(debug=True)
    except Exception as e:
        logger.error(f"Error starting application: {str(e)}")
        raise e 